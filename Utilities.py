import os
from os import path
import time
from datetime import datetime
import pandas as pd

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from selenium.webdriver.support.select import Select
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import *

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import re
import glob
import random
import string
import Logs as l

def lancer_SISTER():
    global driver
    driver = webdriver.Chrome("C:\webdrivers\chromedriver.exe")
    driver.get("https://x-x.x.x.xx.com")
    #driver.get("https://x-x.x.x.xx.com")
    # driver.find_element_by_xpath("//*[@name='userid']").send_keys('x')
    driver.maximize_window()

def cliquer1(xpath):
    element = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, xpath)))
    driver.execute_script("arguments[0].click();", element)

def scroll(xpath):
    element = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, xpath)))
    driver.execute_script("arguments[0].scrollIntoView();", element)

def remplir_champ(xpath, valeur):
    champ = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, xpath)))
    if champ != '':
        champ.clear()
    elif champ == '':
        pass
    champ.send_keys(valeur)

class cliquer:
    def __init__(self, nom1):
        self.nom1 = nom1
        self.action1 = WebDriverWait(driver, 25).until(EC.presence_of_element_located((By.XPATH, nom1)))
        time.sleep(1)
        self.action1.click()

    def __init__(self, nom2):
        self.nom2 = nom2
        self.action2 = WebDriverWait(driver, 25).until(EC.element_to_be_clickable((By.XPATH, nom2)))
        time.sleep(1)
        self.action2.click()

def press_enter():
    try:
        ActionChains(driver).send_keys(Keys.ENTER).perform()
    except Exception as e:
        ActionChains(driver).send_keys(Keys.ENTER).perform()

def rechercher_EAN13(articlenum, desc, org):
    cliquer("//a[@title='Navigateur']")
    cliquer("//div[@title = 'Gestion des produits']")
    cliquer("//a[@title='Gestion des informations produit']")  # ---> (2)
    cliquer("//div[@title='Tâches']")
    cliquer("//*[text()='Gérer les articles']")
    cliquer("//button[text()='Rechercher']")  # ---> (7)

    time.sleep(2)
    remplir_champ("//input[contains(@id, ':value00::content')]", articlenum)
    remplir_champ("//input[contains(@id, ':value10::content')]", desc)
    remplir_champ("//input[contains(@id, ':value50::content')]", org)
    time.sleep(7)
    elem = driver.find_element_by_xpath("//button[text() = 'Rechercher']") # 227481
    actions = ActionChains(driver)
    actions.move_to_element(elem).click().perform()

def val_random(ddown, xpathoptions, xpathind, label, option1vide):
    cliquer(ddown)
    Nombre = len(wait.until(EC.presence_of_all_elements_located((By.XPATH, xpathoptions))))
    print("Nombre d'elements pour " + label + " est:", Nombre)  # Expected output est correct
    print("Les options pour " + label + " sont:")
    try:
        if option1vide == False:
            for i in range(1, Nombre + 1):
                scroll(xpathind + str(Nombre) + "]")
                ele = wait.until(EC.presence_of_element_located((By.XPATH, xpathind + str(i) + "]")))
                print(ele.text)  # to find all elements
                time.sleep(0.1)
            randindex1 = random.randint(1, Nombre)
            element = wait.until(EC.presence_of_element_located((By.XPATH, xpathind + str(randindex1) + "]")))
            time.sleep(0.1)
            driver.execute_script("arguments[0].click();", element)  # to select a random value
        elif option1vide == True:
            for i in range(2, Nombre + 2):
                scroll(xpathind + str(Nombre + 1) + "]")
                ele = wait.until(EC.presence_of_element_located((By.XPATH, xpathind + str(i) + "]")))
                print(ele.text)
                time.sleep(0.1)
            randindex2 = random.randint(2, Nombre + 1)
            element = wait.until(EC.presence_of_element_located((By.XPATH, xpathind + str(randindex2) + "]")))
            time.sleep(0.1)
            driver.execute_script("arguments[0].click();", element)  # to select a random value
    finally:
        print()
    time.sleep(2)

valeur_M = ['BrandAT_Display', 'CommercialModelAT', 'ColorAT', 'ModelCodeAT', 'MobileNaCodeAT', 'LabelModelAT', 'ReferenceTypeAT', 'SavEligibilityAT', 'CPSAT', 'CDRAT', 'EANUsageAT',\
             'SupplierRefBoolAT', 'SupplierAT_Display', 'PmsEligibilityAT', 'UsingArticleInDemoAT', 'SupplierApproDemoAT', 'TechnologyAT', 'MemoryCapacityAT', 'TraceabilityAT',\
             'ExpTraceabilityCodeAT', 'StartDateAT'] # Pour trouver les valeurs des champs

#label_M = ['Marque', 'Modèle C', 'Couleur', 'code', 'NA', 'Libellé', 'Type R', 'SAV', 'CPS', 'CDR', 'EAN13', 'Référencement', 'Fournisseur', 'PMS', 'demo', 'Démo a', 'Technologie', 'Mémoire',\
#            'Niveau', 'attendue', 'initialisation'] # Pour trouver les Labels

label_Dir = ['Cluster']
valeur_Dir = ['ClusterAT']

dict_M = {'Marque': 'BrandAT_Display', 'Modèle C': 'CommercialModelAT', 'Couleur': 'ColorAT', 'code': 'ModelCodeAT', 'NA': 'MobileNaCodeAT', 'Libellé':'LabelModelAT', 'Type R': 'ReferenceTypeAT',\
          'SAV': 'SavEligibilityAT', 'CPS': 'CPSAT', 'CDR': 'CDRAT', 'Utilisation EAN13': 'EANUsageAT', 'EAN 13 Article': 'EANItemAT', 'Référencement fo': 'SupplierRefBoolAT',\
          'PMS': 'PmsEligibilityAT', 'demo': 'UsingArticleInDemoAT', 'Démo a': 'SupplierApproDemoAT', 'Technologie': 'TechnologyAT', 'Mémoire': 'MemoryCapacityAT', 'Niveau': 'TraceabilityAT',\
          'attendue': 'ExpTraceabilityCodeAT', 'initialisation': 'StartDateAT'} # Pour trouver les Labels
# 'Fournisseur': 'SupplierAT_Display'
dict_Dir = {'Cluster': 'ClusterAT'}

dict_A = {'Producteur': 'EcotaxProducerAT', 'achat HT': 'PurchasingPriceAT', 'unitaire RCP HT (€)': 'RcpAT', 'Ecotaxe HT': 'EcotaxAT', 'implanté en France': 'FraImplantSupplierAT'}

dict_GP = {'Donneur d': 'ContractorAT', 'ligne produit 1': 'SegmentProdLine1AT', 'ligne produit 2': 'SegmentProdLine2AT', 'ligne produit 3': 'SegmentProdLine3AT', 'ligne produit 4': 'SegmentProdLine4AT',\
           'Item Family': 'ItemFamilyAT', 'Item Commodity Code': 'ItemCommodityCodeAT', 'expédition PDL': 'PcbAftAT', 'Enjeu Commercial': 'CommercialStakeAT'}

dict_FSC = {'Famille ASTR': 'AstrAT', 'Famille PHENIX': 'PhenixAT', 'Famille SIANT': 'SiantAT', 'Famille de Gestion': 'FamilyManagementAT'}
#print(test())

def creation_dict(DICT):
    dict_val = {}
    for label_xp, val_xp in DICT.items():
        scroll("//label[contains(text(), '" + label_xp + "') and contains(@class, 'label-text')]")
        label = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//label[contains(text(), '" + label_xp + "') and contains(@class, 'label-text')]"))).text
        val = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, '" + val_xp + "')]"))).get_attribute('value')
        dict_val[label] = val
    return dict_val
    
def creation_liste(xpath_labels, xpath_values):  # fonction pour recuperer les valeurs et les presentes
    wait = WebDriverWait(driver, 30)
    global list_valeurs, list_label
    list_label = []
    list_valeurs = []

    for value in xpath_labels:
        label = wait.until(EC.presence_of_element_located((By.XPATH, "//label[contains(text(), '" + value + "') and contains(@class, 'label-text')]"))).text
        list_label.append(label)
    #list_label.pop(19)
    #list_label.insert(19, 'Fournisseur') # a la fin
    for value in xpath_values:
        valeur = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, '" + value + "')]"))).get_attribute('value')
        list_valeurs.append(valeur)
    return list_valeurs

# Exceptionellement, que pour les premieres valeurs a collecter, ecrire les codes dans le script main.py
# Ensuite, creer une fonction uniforme pour pouvoir append les valeurs
# Finalement, creer une fonction pour s'en servir de la liste complete pour les exporter sur l'Excel

def deconnexion():
    target = wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@title = 'Menu Paramètres et actions']")))
    actions = ActionChains(driver)
    actions.move_to_element(target).click().perform()
    element = wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@title = 'Menu Paramètres et actions']")))
    driver.execute_script("arguments[0].click();", element)
    cliquer("//*[text()='Déconnexion']")
    cliquer("//button[@name='Confirm']")

def login_1(identifiant, password= 'xxx'):
    userid = wait.until(EC.presence_of_element_located((By.NAME, 'userid')))
    userid.send_keys(identifiant)
    driver.find_element_by_xpath("//*[@name='password']").send_keys(password)
    cliquer("//button[text() = 'Connexion ']")

def notif(num1, num2):
    #    i += 1
    #    #ic = WebDriverWait(driver, 15).until(EC.visibility_of_element_located((By.XPATH, "//*[@id='pt1:_UISatr:0:cil1::icon']")))
    #    #driver.execute_script("arguments[0].click();", ic)
    #    try:
    #        actions.move_to_element(target).click().perform()
    #        time.sleep(2)
    #    except StaleElementReferenceException as e:
    #        Cliquer("//*[@id='pt1:_UISatr:0:cil1::icon']")
    #        time.sleep(2)
    i = 1
    while i <= 5:
        print("New i", i)
        #i += 1
        target = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//*[@id='pt1:_UISatr:0:cil1::icon']")))
        actions = ActionChains(driver)
        actions.move_to_element(target).click().perform()
        try:
            for k in [0, 5]:
                ele = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//div[@data-afrrk = '"+str(k)+"']//td/span[contains(@title, 'secondes')]")))
                #ele1 = WebDriverWait(driver, 7).until(EC.presence_of_element_located((By.XPATH, "//div[@data-afrrk = '5']//td/span[contains(@title, 'secondes')]")))
                ele1 = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//div[@data-afrrk = '"+str(k)+"']//td/a[contains(@title, 'New Item Request NIR-')]")))
                if ele.is_displayed() and ele1.is_displayed():
                    print(ele.text)
                    print(ele1.text)
                    print("ok", k)
                    print(i)
                    target1 = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//div[@data-afrrk = '" + str(k) + "']//td/a[contains(@title, 'New Item Request NIR-')]")))
                    actions = ActionChains(driver)
                    actions.move_to_element(target1).click().perform()
                    if i < 5:
                        i += 5
                        print("Here", i)
                        break
                elif not(ele.is_displayed() and ele1.is_displayed()):
                    print("Pas trouver !", i, "\n")
                    driver.refresh()
                    i += 1
        except Exception as e:
            print(e)
    # icon = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//*[@id='pt1:_UISatr:0:cil1::icon']")))
    #icon = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//*[@id='pt1:_UISatr:0:cil1::icon']")))
    #actions = ActionChains(driver)
    # driver.execute_script("arguments[0].click();", icon)
    #actions.move_to_element(icon).click().perform()

    #"//td[@id='pt1:_UISatr:0:lv4:0:gc6']/a[contains(@title,'La demande de nouvel article NIR-')]"
    #"//div[@data-afrrk = '0']//td/span[contains(@title, 'secondes')]//td/a[contains(@title, 'New Item Request NIR-')]"
    # "//tr/td[2]/a[contains(@title, 'New Item Request NIR-')]"
    #Cliquer("//div[@data-afrrk = '0']//tr[2]/td[2]") !!!!!!!!!!!!!!!!!!!!!!!!!
    time.sleep(7)
    window_new1 = driver.window_handles[num1]
    driver.switch_to.window(window_new1)

    NIR = wait.until(EC.presence_of_element_located((By.XPATH, "//a[@class = 'xij' and contains(text(), 'NIR')]"))).text
    print(NIR)
    cliquer("//*[@title='Aller à la tâche']")
    time.sleep(7)
    window_recent = driver.window_handles[num2]
    driver.switch_to.window(window_recent)
    WebDriverWait(driver, 70).until(EC.presence_of_element_located((By.XPATH, "//a[text()='Spécifications']"))).click()

class fonctionnel:
    def __init__(self, wait, driver, mdp, titre, nom, img_chemin, valeur, xpath, nom_repertoire, lot,
                 tnr_69, num_tnr, utilisateur):
        self.wait = wait
        self.driver = driver
        self.mdp = mdp
        self.titre = titre
        self.img_chemin = img_chemin
        self.nom = nom
        self.valeur = valeur
        self.xpath = xpath
        self.nom_repertoire = nom_repertoire
        self.lot = lot
        self.tnr_69 = tnr_69
        self.num_tnr = num_tnr
        self.utilisateur = utilisateur

    def creation_dossier(nom_repertoire):
        #directory = 'TNR-94'
        global Directory, parent_dir
        parent_dir = "C:/Users/" + os.getlogin() + "/Desktop/"
        Directory = os.path.join(parent_dir, nom_repertoire)
        if os.path.exists(Directory):
           print('Folder exist')
        else:
           print('Folder not found')
           os.mkdir(Directory)
           print("Directory '% s' created" % Directory)
        return Directory

    def capture_ecran_doc(nom):
        global doc, fname_w
        fname_w = Directory + str(nom) + datetime.today().strftime('%Y-%m-%d') + "_" + time.strftime("%H.%M.%S", time.localtime()) + ".docx" # /tnr90_
        #le word doit etre en dehors de la fonction
        if os.path.exists(fname_w) == True:
           print("The filename for the word document is outdated. Please update!")
           breakpoint()
        elif os.path.exists(fname_w) == False:
           doc = docx.Document()

    def login(self, num_tnr):
        global driver, wait
        driver = webdriver.Chrome("C:\webdrivers\chromedriver.exe")
        wait = WebDriverWait(driver, 25)
        driver.get("https://x-x.x.x.xx.com")
        driver.maximize_window()

        language = wait.until(EC.presence_of_element_located((By.NAME, 'Languages')))
        if language != "French - français":
            language.click()
            fr = wait.until(EC.element_to_be_clickable((By.XPATH, "//option[@value='fr-fr']")))
            fr.click()
        else:
            pass
        userID = wait.until(EC.presence_of_element_located((By.NAME, 'userid')))
        if num_tnr == '87':
            userID.send_keys('CDG')
        ### !!! elif num_tnr == '94' or '95' or 'CA':
        ### !!!     userID.send_keys('XXSCF_MOA')
        #elif num_tnr == 'CA':
        #    userID.send_keys('RGA')
        elif num_tnr == '94' or num_tnr == '95':
            userID.send_keys('XXSCF_MOA')
        else:
            userID.send_keys('RGA')
        driver.find_element_by_xpath("//*[@name='password']").send_keys('xxx!')
        connex = wait.until(EC.presence_of_element_located((By.XPATH, "//button[@id = 'btnActive']")))
        connex.click()

    def demarrer_tnrSISTER(self, nom_repertoire, nom, num_tnr):
        fonctionnel.creation_dossier(nom_repertoire)
        fonctionnel.capture_ecran_doc(nom)
        fonctionnel.login(self, num_tnr)

    def vers_Word(titre, img_chemin):
       # doc = docx.Document()
        doc.add_paragraph(titre)
        doc.add_picture(img_chemin, width=docx.shared.Inches(7.2), height=docx.shared.Inches(4))
        doc.save(fname_w)

    def vers_Excel(fname, list_label, list_valeurs):
        # fname = "C:/Users/xxx/Desktop/Achats/test.csv"
        FileExist = 1
        if not path.exists(fname):
            FileExist = 0
        with open(fname, 'a') as filehandle:
            if FileExist == 0:
                for lab in list_label:
                    filehandle.write('%s,' % lab)
                filehandle.write('\n')
                for val in list_valeurs:
                    filehandle.write('%s,' % val)
                filehandle.write('\n')
            elif FileExist == 1:
                for val in list_valeurs:
                    filehandle.write('%s,' % val)
                filehandle.write('\n')

    def chercher_creation_article_mobile(self):
        cliquer("//div[@title='Tâches']")
        cliquer("//*[contains(text(), 'Gérer les lots d')]")
        Nom = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'pt1:AP1:r1:0:q1:value10::content')]")))
        Nom.send_keys('creation')

        #if no_result.text() == 'Aucun résultat trouvé.':
        #while False:
        #    print("Aucun lot n'est apparu")
        #    driver.get_screenshot_as_file("C:/Users/"+os.getlogin()+"/Desktop/TNR-"+str(num_art)+"/C_A_M_G.png")
        #    fonctionnel.vers_Word("Aucun lot d'article disponible. La réaffectation est nécessaire", "C:/Users/"+os.getlogin()+"/Desktop/TNR-"+str(num_art)+"/C_A_M_G.png")
        #    try:
        #        cliquer("//select[@title = 'Egal ou postérieur à']")  # --> (1) <--
        #        cliquer("//button[text() = 'Rechercher']")
        #        # num = U.WebDriverWait(U.driver, 20).until(U.EC.presence_of_element_located((U.By.XPATH, "//span[text() = '22007']")))
        #        # num.click()
        #        cliquer("//span[text() = '22007']")
        #        time.sleep(3)
        #        reaffecter()
        #    except TimeoutException as e:
        #        cliquer("//select[@title = 'Egal ou postérieur à']")
        #        cliquer("//option[@value = 'ONORBEFORE' and @title = 'Egal ou antérieur à']")  # --> (2) <--
        #        cliquer("//button[text() = 'Rechercher']")
        #        cliquer("//span[text() = '22007']")
        #        reaffecter()

        #except TimeoutException as e:
        #    print("Le lot d'article est utilisable. Pas besoin de faire la reaffectation")
        #    driver.get_screenshot_as_file("C:/Users/"+os.getlogin()+"/Desktop/TNR-"+str(num_art)+"/C_A_M.png")
        #    fonctionnel.vers_Word("Le lot d'article est utilisable. Nul besoin de faire la reaffectation.", "C:/Users/"+os.getlogin()+"/Desktop/TNR-"+str(num_art)+"/C_A_M.png")
        #    #Action("//td[contains(text(), 'Modifier les options du lot d')]")
        #    cliquer("//span[text() = '22007']")
        #    cliquer1("//div[@aria-label = 'Actions']")
        #    time.sleep(3)
        #    cliquer1("//td[contains(text(), 'Modifier les options du lot d')]")
        #    pass

    #def Tx_Depre_initial(self, num_tnr, num_art):

    def statut_du_chargement_de_fichier(self, num_tnr):
        for i in range(25):
            statut = WebDriverWait(driver, 8).until(EC.presence_of_element_located((By.XPATH, "//span[contains(@id, ':SearchResultsTable:_ATp:blt:0:soc1::content')]")))
            if statut.text == 'En attente':
                driver.refresh()
                time.sleep(2)
                print('En attente')
            elif statut.text == 'Actif':
                print('Resultat connu')
                driver.get_screenshot_as_file("C:/Users/xxx/Desktop/TNR-/"+str(num_tnr)+"/result.png")
                fonctionnel.vers_Word("Statut de l'import connu !", "C:/Users/xxx/Desktop/TNR-"+str(num_tnr)+"/result.png")
                time.sleep(3)
                try:
                    driver.find_element_by_xpath("//img[contains(@id, ':_ATp:blt:0:cil45::icon')]").click()
                except TimeoutException as e:
                    driver.find_element_by_xpath("//img[contains(@id, ':_ATp:blt:0:cil45::icon')]").click()
                statut1 = WebDriverWait(driver, 7).until(EC.presence_of_element_located((By.XPATH, "//a[contains(@id, ':j_id__ctru66pc8')]")))
                if statut1.text == 'Erreur':
                    print('Erreur')
                    cliquer("//button[contains(@id, 'AP1:r1:0:SearchResultsTable:d8::ok')]")
                    # U.Cliquer("//button[contains(@id, ':pt1:AP1:cb7')]") # Terminer
                    cliquer("//a[@title = 'RGA' and contains(@id, 'pt1:_UIScmil2u')]")
                    cliquer("//a[text() = 'Déconnexion']")
                    cliquer("//button[@name = 'Confirm']")
                    userid = wait.until(EC.presence_of_element_located((By.NAME, 'userid')))
                    userid.send_keys("RGA")
                    driver.find_element_by_xpath("//*[@name='password']").send_keys('xxx')
                    sign = wait.until(EC.presence_of_element_located((By.XPATH, "//button[@id = 'btnActive']")))
                    sign.click()
                    break
                elif statut1.text == 'Réussite':
                    driver.get_screenshot_as_file("C:/Users/xxx/Desktop/TNR-"+str(num_tnr)+"/import_reussi.png")
                    fonctionnel.vers_Word("L'import a abouti !", "C:/Users/SXBS5132/Desktop/TNR-"+str(num_tnr)+"/import_reussi.png")
                    cliquer("//button[contains(@id, 'AP1:r1:0:SearchResultsTable:d8::ok')]")
                    cliquer("//img[@title = 'Actualiser']")
                    break              
'''
Reaffectation lot d'article tnr69:
1. 
'''
class fn_tnr25:
    def __init__(self, classe_article):
        self.article = classe_article

    def renseigner_classe_article(self, classe_article):
        try:
            cliquer("//a[@title = 'Navigateur']")
        except TimeoutException as e:
            driver.refresh()
            cliquer("//a[@title = 'Navigateur']")
        cliquer("//div[@title='Gestion des produits']")
        cliquer("//a[@title='Gestion des informations produit']")
        cliquer("//div[@title='Tâches']")
        cliquer("//*[text()='Créer un article']")
        cliquer("//a[contains(@title, 'Recherche : Classe d')]")
        id_liste_articles = wait.until(EC.presence_of_element_located((By.ID, '_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:itemClassId::dropdownPopup::dropDownContent::scroller')))
        for i in range(1, 100):
            try:
                article = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.XPATH, "//div[contains(@id, 'ClassId::dropdownPopup::dropDownContent')]//table//tr[" + str(i) + "]/td[2]")))

                if article.text != classe_article:  # <---------------------------
                    print(article.text)
                    id_liste_articles.send_keys(Keys.ARROW_DOWN)
                elif article.text == classe_article:
                    print("L'article en question est: ", classe_article)
                    cliquer("//td[1]/span[text() = '" + str(classe_article) + "']")  # <------------------------------------------------------
                    break
            except TimeoutException as e:
                break
                
        cliquer("//*[text()='RECYCLE']")  # !!!!!!!!!!
        cliquer("//a[@title='Déplacez les éléments sélectionnés vers : Liste de sélection']")
        cliquer("//*[text()='Fini']")
        cliquer("//a[@title='Déplacez les éléments sélectionnés vers : Liste de sélection']")
        cliquer("//button[@accesskey='K']")
        driver.save_screenshot(Directory + "/multimedia.png")
        fonctionnel.vers_Word("Créer un article MULTIMEDIA  en matrice « Fini » et référence « RECYCLE »", Directory + "/multimedia.png")

        cliquer("//a[text()='Spécifications']")
        cliquer("//a[text()='SISTER: Marketing']")

    def renseigner_Marketing(self):
        global result, dic1, Fourni_label, Fourni_valeur
        # -------> Description Communication
        val_random("//span[contains(@id, 'BrandAT_Display::cntnrSpan')]", \
                     "//div[contains(@id, 'BrandAT_Display::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'BrandAT_Display::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Code Marque', False)  # Code Marque

        letters = string.ascii_letters  # Modèle Commercial
        result = ''.join(random.choice(letters) for i in range(4))
        remplir_champ("//input[contains(@id, '_ItemXxscfCustomerProductsIcPrivateVOXXSCF_COMMERCIAL_DESCRIPTION_AG:0:XXSCFCommercialModelAT')]", 'Test' + result)
        # pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sspecf:1:dynReg:1:CTXRNj_ItemXxscfBureautiqueIcPrivateVOXXSCF_COMMERCIAL_DESCRIPTION_AG:0:XXSCFCommercialModelAT pour BUREAUTIQUE
        val_random("//span[contains(@id, 'ColorAT::cntnrSpan')]", \
                     "//div[contains(@id, 'ColorAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'ColorAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Couleur', False)  # Couleur

        remplir_champ("//input[contains(@id, '_ItemXxscfCustomerProductsIcPrivateVOXXSCF_COMMERCIAL_DESCRIPTION_AG:0:XXSCFModelCodeAT')]", 'HUA P30')

        # -------------> Missing: NA Code, Libelle Modele
        # --------------> Informations générales
        val_random("//span[contains(@id, 'SavEligibilityAT::cntnrSpan')]", \
                     "//div[contains(@id, 'SavEligibilityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'SavEligibilityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Eligibilité SAV', False)  # Eligibilité SAV

        val_random("//span[contains(@id, 'CPSAT::cntnrSpan')]", \
                     "//div[contains(@id, 'CPSAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'CPSAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'CPS - Catégorie produit et service', False)  # CPS - Catégorie produit et service

        val_random("//span[contains(@id, 'CDRAT::cntnrSpan')]", \
                     "//div[contains(@id, 'CDRAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'CDRAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'CDR - Centre de responsabilité', False)  # CDR - Centre de responsabilité

        # Génération EAN
        val_random("//span[contains(@id, 'EANUsageAT::cntnrSpan')]", \
                     "//div[contains(@id, 'EANUsageAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'EANUsageAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Utilisation EAN13 Fournisseur', False)  # Utilisation EAN13 Fournisseur

        # Référence fournisseur
        val_random("//span[contains(@id, 'SupplierRefBoolAT::cntnrSpan')]", \
                     "//div[contains(@id, 'SupplierRefBoolAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[2][string-length()>1]", \
                     "//div[contains(@id, 'SupplierRefBoolAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Référencement fournisseur', True)  # Référencement fournisseur

        scroll("//span[contains(@id, 'EANUsageAT::cntnrSpan')]")

        val_random("//span[contains(@id, 'SupplierAT_Display::cntnrSpan')]", \
                     "//div[contains(@id, 'SupplierAT_Display::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[2][string-length()>1]", \
                     "//div[contains(@id, 'SupplierAT_Display::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Fournisseur', True)  # Fournisseur

        # Caractéristiques Produit
        val_random("//span[contains(@id, 'PmsEligibilityAT::cntnrSpan')]", \
                     "//div[contains(@id, 'PmsEligibilityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'PmsEligibilityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Eligibilité PMS', False)  # Eligibilité PMS

        UAeD = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, 'PRODUCT_CHARACTERISTICS_AG:0:XXSCFUsingArticleInDemoAT')]")))
        UAeD.clear()
        val_random("//span[contains(@id, 'UsingArticleInDemoAT::cntnrSpan')]", \
                     "//div[contains(@id, 'UsingArticleInDemoAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][string-length(text())=0]", \
                     "//div[contains(@id, 'UsingArticleInDemoAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Utilisation article en demo', True)  # Utilisation de l'article en demo

        try:
            time.sleep(2)
            UAeD_value = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, 'PRODUCT_CHARACTERISTICS_AG:0:XXSCFUsingArticleInDemoAT')]"))).get_attribute('value')
            print("LA VALEUR de l'utilisation de l'article en demo EST: ", UAeD_value)
            print()
            if UAeD_value == 'O':  # Démo appro fournisseur obligatoire si UAeD est 'OUI'
                val_random("//span[contains(@id, 'PRODUCT_CHARACTERISTICS_AG:0:XXSCFSupplierApproDemoAT::cntnrSpan')]", \
                    "//div[contains(@id, 'PRODUCT_CHARACTERISTICS_AG:0:XXSCFSupplierApproDemoAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[2][string-length()>1]", \
                    "//div[contains(@id, 'PRODUCT_CHARACTERISTICS_AG:0:XXSCFSupplierApproDemoAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                    'Démo appro fournisseur', True)  # Démo appro fournisseur
            elif UAeD_value == 'N':
                pass
        except Exception as e:
            pass

        val_random("//span[contains(@id, 'TechnologyAT::cntnrSpan')]", \
                     "//div[contains(@id, 'TechnologyAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][string-length()>1]", \
                     "//div[contains(@id, 'TechnologyAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Technologie', True)  # Technologie

        val_random("//span[contains(@id, 'MemoryCapacityAT::cntnrSpan')]", \
                     "//div[contains(@id, 'MemoryCapacityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][span[string-length(text())>0]]", \
                     "//div[contains(@id, 'MemoryCapacityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Capacité Mémoire (Go)', True)  # Capacité Mémoire (Go)

        # Traçabilité
        val_random("//span[contains(@id, 'TraceabilityAT::cntnrSpan')]", \
                     "//div[contains(@id, 'TraceabilityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'TraceabilityAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Niveau de traçabilité', False)  # Niveau de traçabilité

        scroll("//span[contains(@id, 'MemoryCapacityAT::cntnrSpan')]")
        val_random("//span[contains(@id, 'ExpTraceabilityCodeAT::cntnrSpan')]", \
                     "//div[contains(@id, 'ExpTraceabilityCodeAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr", \
                     "//div[contains(@id, 'ExpTraceabilityCodeAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Traçabilité attendue', False)  # Traçabilité attendue

        scroll("//*[text() = 'Accueil']")

        cliquer("//span[text()='Sauvegarder']")
        try:
            WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//td[text() = 'Erreur: La valeur est obligatoire.']")))

            while True:
                driver.save_screenshot(Directory + "/save.png")
                fonctionnel.vers_Word("Sauveagarde echec du au champs description non remplie automatiquement",
                                      Directory + "/save.png")
                break
                ### !!! print("Il faut renseigner la description qui etait auparavant faite automatiquement")
                ### !!! descr = driver.find_element_by_xpath("//textarea[contains(@name, 'MAt3:0:pt1:ap1:r10:0:inputText2')]")
                ### !!! descr.send_keys('Test' + result)
                ### !!! break
        finally:
            cliquer("//span[text()='Sauvegarder']")
            time.sleep(15)

        # "//div[contains(text(), 'Merci de renseigner le code EAN (EAN fournisseur est à Oui)')]"
        fn_tnr25.trySaveSiEAN_O(self)  # EAN fournisseur doit prendre la valeur N
        # "//div[contains(text(), 'Démo appro fournisseur obligatoire')]"
        ###try:
        ###    # time.sleep(2) # "//div[@class = 'x1ii']/a[contains(@class, 'p_AFSelected') and text() = 'Spécifications']" # "//div[@class = 'x1ii']/a[text() = 'Spécifications']"
        ###    speci = wait.until(EC.element_to_be_clickable((By.XPATH, "//div[contains(@_afrptkey, 'sspec')]")))
        ###    ActionChains(driver).click(speci).perform()
        ###    time.sleep(3) # "//div[@class = 'xb7']//a[text()='Spécifications']"
        ###except TimeoutException as e:
        ###    speci = WebDriverWait(driver, 15).until(EC.element_to_be_clickable((By.XPATH, "//div[contains(@_afrptkey, 'sspec')]")))
        ###    ActionChains(driver).click(speci).perform()
        ###    time.sleep(3)
        ###finally:
        ###    time.sleep(2)
        ###    cliquer("//a[text()='SISTER: Marketing']")
###
        ###dic1 = creation_dict(dict_M)
        ###print(dic1)
        ###scroll("//tr[contains(@id, 'SupplierAT_Display')]/td/label[text() = 'Fournisseur']")
        ###Fourni_label = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//tr[contains(@id, 'SupplierAT_Display')]/td/label[text() = 'Fournisseur']"))).text
        ###Fourni_valeur = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, 'SupplierAT_Display')]"))).get_attribute('value')

    def trySaveSiEAN_O(self):  # Void
        # targetEAN = WebDriverWait(driver, 15).\
        #    until(EC.presence_of_element_located((By.NAME, 'pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sspecf:1:dynReg:1:CTXRNj_ItemXxscfOrangeIcPrivateVOXXSCF_EAN_GENERATION_AG:0:XXSCFEANUsageAT'))).get_attribute('value')

        # wait.until(lambda driver: driver.find_element_by_xpath(xpathApercu)) # Par defaut, l'onglet Specifications est deja selecte. Il faut attendre que l'onglet Apercu soit selecte automatiquement.

        while True:
            ok = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//button[@_afrpdo='cancel' and text()='OK']")))
            if ok.text == 'OK':
                print("Sauvegarde n'a pas eu lieu. Merci de renseigner le code EAN (EAN fournisseur est à Oui)")
                cliquer("//button[@_afrpdo='cancel' and text()='OK']")
                scroll("//*[contains(@name, 'EAN_GENERATION_AG:0:XXSCFEANUsageAT')]")
                print("test123")
                driver.find_element_by_xpath("//input[contains(@name, '_ItemXxscfOrangeIcPrivateVOXXSCF_EAN_GENERATION_AG:0:XXSCFEANUsageAT')]").clear()
                driver.find_element_by_xpath("//input[contains(@name, '_ItemXxscfOrangeIcPrivateVOXXSCF_EAN_GENERATION_AG:0:XXSCFEANUsageAT')]").send_keys('N')
                print("EAN fournisseur est maintenant NON")
                scroll("//*[text() = 'Accueil']")
                cliquer("//span[text()='Sauvegarder']")
                time.sleep(8)
                driver.refresh()
                time.sleep(5)
                break
            else:
                print("Sauvegarde reussi. EAN est a NON deja.")
                driver.refresh()
                time.sleep(5)
                break

    def renseigner_Dirco(self):
        scroll("//*[@title = 'Accueil']")
        cliquer("//div[@_afrptkey = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec']")

        try:
            cliquer("//a[text()='SISTER: Direction commerciale']")
        except ElementClickInterceptedException as e:
            Dirco = driver.find_element_by_xpath("//a[text()='SISTER: Direction commerciale']")
            driver.execute_script("arguments[0].click();", Dirco)

        val_random("//span[contains(@id, 'CLUSTER_AG:0:XXSCFClusterAT::cntnrSpan')]", \
                     "//div[contains(@id, 'ClusterAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][span[string-length(text())>0]]", \
                     "//div[contains(@id, 'ClusterAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Cluster', True)  # Cluster

        dic2 = creation_dict(dict_Dir)
        dic3 = {**dic1, **dic2}
        print(dic3)

        lab = dic3.keys()
        lab_list = list(lab)
        lab_list.insert(13, Fourni_label)
        #
        val = dic3.values()
        val_list = list(val)
        val_list.insert(13, Fourni_valeur)

        dic4 = dict(zip(lab_list, val_list))
        print("dic4: ", dic4)

        # U.vers_Excel(dic3)
        # U.vers_Excel(lab_list, val_list) # creer une fonction pour "//tr[contains(@id, 'SupplierAT_Display')]/td/label[text() = 'Fournisseur']" pour ajouter a la liste a chaque fois

    def renseigner_Achat(self):
        # ACHAT
        #scroll("//a[text()='Spécifications']")
        #try:
        #    cliquer("//a[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec::disAcr']")
        #except ElementClickInterceptedException as e:
        #    print("ElementClickInterceptedException")
        #    for i in range(8):
        #        Spec = driver.find_element_by_xpath("//a[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec::disAcr']")
        #        driver.execute_script("arguments[0].click();", Spec)
        #try:
        #    # 1. xpathApercu = "//a[text()='Aperçu' and @id='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sover::disAcr' and contains(@class, ' p_AFSelected')]" # ' p_AFSelected' n'y est plus dans le HTML
        #    # xpathApercu = "//div[@_afrptkey='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sover']/div/a[text()='Aperçu' and contains(@id, ':sover::disAcr')]"
        #    # wait.until(lambda driver: driver.find_element_by_xpath(xpathApercu)) # Par defaut, l'onglet Specifications est deja selecte. Il faut attendre que l'onglet Apercu soit selecte automatiquement.
        #    cliquer("//a[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec::disAcr']")
        #    # spec = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//div[@id='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sspec::ti']//div[a[text()='Spécifications']]")))
        #    # spec.click()
        #except TimeoutException as e:
        #    # Cliquer("//div[@id='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sspec::ti']//div[a[text()='Spécifications']]")
        #    # element = wait.until(EC.presence_of_element_located((By.XPATH, "//div[@id='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sspec::ti']//div[a[text()='Spécifications']]")))
        #    element = driver.find_element_by_xpath("//a[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec::disAcr']")
        #    driver.execute_script("arguments[0].click();", element)
        #    driver.execute_script("arguments[0].click();", element)
        #    time.sleep(3)

        #scroll("//a[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec::disAcr']")

        Spec = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//a[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspec::disAcr']")))
        if Spec.text == 'Spécifications':
            print("Spec trouver", Spec.text)
        else:
            print("NON spec")

        try:
            Spec.click()
        except Exception as e:
            time.sleep(3)
            driver.execute_script("arguments[0].click();", Spec)
            
        time.sleep(5)
        try:
            WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//div[text() = 'Messages des règles']")))
            while True:
                driver.get_screenshot_as_file("C:/Users/SXBS5132/Desktop/Achats/Anomalie" + result + ".png")
                cliquer("//td[@class = 'p_AFResizable x1rh']/button[contains(@id, 'ap1:d4') and text() = 'OK']")
                break
            else:
                pass
        except TimeoutException as e:
            pass
        time.sleep(5)
        achat = driver.find_element_by_xpath("//a[text()='SISTER: Achat']")
        driver.execute_script("arguments[0].click();", achat)
        time.sleep(5)
        driver.execute_script("arguments[0].click();", achat)

        scroll("//label[text()='Producteur' and @title=' Orange Producteur Ecotaxe']")
        val_random("//span[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFEcotaxProducerAT::cntnrSpan')]", \
                     "//div[contains(@id, 'EcotaxProducerAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][span[string-length(text())>0]]", \
                     "//div[contains(@id, 'EcotaxProducerAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Producteur', True)  # Producteur doit prendre la valeur oui

        PrixAchat = wait.until(EC.presence_of_element_located((By.XPATH,
                    "//input[contains(@name, 'PURCHASE_INFORMATION_AG:0:XXSCFPurchasingPriceAT')]")))  # Le nom varie !!!
        PrixAchat.send_keys('25')

        val_random("//span[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFRcpAT::cntnrSpan')]", \
                     "//div[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFRcpAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][span[string-length(text())>0]]", \
                     "//div[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFRcpAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Montant unitaire RCP HT (€)', True)  # Montant unitaire RCP HT (€)

        val_random("//span[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFEcotaxAT::cntnrSpan')]", \
                     "//div[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFEcotaxAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][span[string-length(text())>0]]", \
                     "//div[contains(@id, 'PURCHASE_INFORMATION_AG:0:XXSCFEcotaxAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Ecotaxe HT (€)', True)  # Ecotaxe HT (€)

        dic5 = creation_dict(dict_A)
        print(dic5)
        #dic6 = {**dic4, **dic5}
        #print("dict6:", dic6)

        cliquer("//span[text()='Sauvegarder']")
        try:
            popUp = WebDriverWait(U.driver, 20).until(EC.presence_of_element_located((By.XPATH, "//div[@id='d1::popup-container' and @data-afr-popupid='d1']")))
            if popUp.is_displayed():
                print("Message d'alerte ! Il faut modifier le champs Producteur et mettre a 'Oui'.")
                cliquer("//div[@id='d1::popup-container' and @data-afr-popupid='d1']//button[@type='button' and text()='OK']")
                Producteur = wait.until(EC.presence_of_element_located((By.XPATH,
                "//input[@aria-describedby='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:sspecf:1:dynReg:2:CTXRNj_ItemXxscfProductsIcPrivateVOXXSCF_PURCHASE_INFORMATION_AG:0:XXSCFEcotaxProducerAT::desc']")))
                Producteur.clear()
                Producteur.send_keys('O')  # Producteur doit prendre la valeur 'Oui'
                # U.scroll("//div[contains(@id, 'pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:customSubmit') and @role='presentation']") # Sauvegarder avant??
                # U.Cliquer("//div[contains(@id, 'pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:customSubmit') and @role='presentation']") #Soumettre
                scroll("//span[text()='Sauvegarder']")
                cliquer("//span[text()='Sauvegarder']")
            else:
                pass
        except Exception as e:
            print("Pas de message d'alerte. Le champs pour Producteur est deja 'Oui'.")

        #try:
        #    def test1():
        #        wait.until(EC.visibility_of_element_located((By.XPATH, "//a[@accesskey='m']")))
        #        cliquer("//a[@accesskey='m']")
        #        time.sleep(3)
        #        cliquer("//a[@accesskey='v']")
#
        #    test1()
        #except Exception as e:
        #    test1()
        #finally:
        #    NOM = U.wait.until(EC.presence_of_element_located((By.XPATH, "//input[@name='pt1:_FOr1:1:_FONSr2:0:MAt3:0:pt1:ap1:SelCh:0:r1:0:it5']")))
        #    ab = 'Test' + result
        #    NOM.send_keys(ab)
        #    cliquer("//span[text()='Sauvegarder et modifier']")
        #    cliquer("//*[text()='Soumettre']")
            deconnexion()

    def test(self):
        login_1('xxx')
        notif(1, 2)

        def changer_fenetre(num):
            window_new1 = driver.window_handles[num]
            driver.switch_to.window(window_new1)

        # changer_fenetre(1)
        # NIR = U.wait.until(EC.presence_of_element_located((By.XPATH, "//a[@class = 'xij' and contains(text(), 'NIR')]"))).text
        # print(NIR)
        # changer_fenetre(2)

        scroll("//*[text() = 'Gestion Produit']")
        val_random("//span[contains(@id, 'XXSCFContractorAT::cntnrSpan')]", \
                     "//div[contains(@id, 'XXSCFContractorAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[2][string-length()>1]", \
                     "//div[contains(@id, 'XXSCFContractorAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Donneur dordre', True)  # Donneur d'ordre

        i = 1
        for i in range(4):
            val_random("//span[contains(@id, 'SegmentProdLine" + str(i + 1) + "AT::cntnrSpan')]", \
                         "//div[contains(@id,'SegmentProdLine" + str(
                             i + 1) + "AT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[1][string-length()>1]", \
                         "//div[contains(@id,'SegmentProdLine" + str(
                             i + 1) + "AT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                         "Segment ligne produit " + str(i + 1), True)

        val_random("//span[contains(@id, 'XXSCFCommercialStakeAT::cntnrSpan')]", \
                     "//div[contains(@id, 'XXSCFCommercialStakeAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[2][string-length()>1]", \
                     "//div[contains(@id, 'XXSCFCommercialStakeAT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                     'Enjeu Commercial', True)  # Enjeu Commercial

        driver.find_element_by_name('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFPcbAftAT').clear()
        # _FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFPcbAftAT
        remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFPcbAftAT',
                       '3')  # Par defaut, la valeur du champs est 1.

        remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFItemFamilyAT',
                       random.randint(1, 9))  # Meme en donnant une valeur aleatoire, on ne peut pas enregistrer.
        # _FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFItemFamilyAT
        cliquer("//*[text()='Enregistrer']")

        WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.XPATH, "//div[text() = 'Avertissement'] | //div[text() = 'Erreur']")))
        while True:
            try:
                # wait.until(lambda driver: driver.find_element_by_xpath(xpathApercu)) # Par defaut, l'onglet Specifications est deja selecte. Il faut attendre que l'onglet Apercu soit selecte automatiquement.
                alert = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH,
                                                                                        "//button[@type = 'button' and text() = 'OK' and contains(@id, '_FOd1::msgDlg::cancel')]")))
                alert.click()
                time.sleep(3)
            except TimeoutException as e:
                ALERT = WebDriverWait(driver, 10).until(EC.element_to_be_clickable(
                    (By.XPATH, "//button[@type = 'button' and text() = 'OK' and contains(@id, 'd5::ok')]")))
                ALERT.click()
                time.sleep(3)
            finally:
                print(
                    "Les valeurs aleatoire pour la section Gestion Produit ne sont pas valides! De ce fait, ces valeurs selon le protocole sont utilises pour remplir les champs.")
                cliquer("//a[@id='_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:sdi9::disAcr' and text()='Spécifications']")
                driver.find_element_by_name(
                    '_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFContractorAT').clear()
                # _FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFContractorAT
                remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFContractorAT', 'DDDXF')
                i = 1
                for i in range(4):
                    driver.find_element_by_name(
                        '_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFSegmentProdLine' + str(
                            i + 1) + 'AT').clear()
                    time.sleep(1)
                    i += 1
                remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFSegmentProdLine1AT',
                               'ACCES')  # name has been changed
                remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFSegmentProdLine2AT', 'MOBIL')
                remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFSegmentProdLine3AT', 'AUTRE')
                remplir_champ('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:r1:1:r15:0:XXSCFSegmentProdLine4AT',
                               'VALO')  # name has been changed
                cliquer("//*[text()='Enregistrer']")
                time.sleep(3)
                try:
                    dic7 = creation_dict(dict_GP)
                    print(dic7)
                except StaleElementReferenceException as e:
                    time.sleep(2)
                    dic7 = creation_dict(dict_GP)
                break
        else:
            pass

        dic8 = {**dic6, **dic7}
        print("dict8: ", dic8)
        # ! FSC Gestion de produits ! ! !
        U.force_click("//div[@class = 'x1ii']/a[text()='Associations']")
        U.force_click("//a[text()='Actions']")
        U.force_click("//td[text()='Sélectionner et ajouter']")

        time.sleep(10)
        U.RemplirChamp('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:AP1:IOF1:0:AT1:qryId1:criterionValue0', 'SCF_OSA_IO_01')
        U.Cliquer("//a[@title='Rechercher']")
        U.Cliquer("//span[text()='SCF_OSA_IO_01']")
        time.sleep(5)
        U.Cliquer("//button[@accesskey='l']")  # Appliquer
        U.Cliquer("//button[@accesskey='e']")  # TerminerU.
        try:
            U.Cliquer("//*[text()='Enregistrer']")
            U.Cliquer("//span[text()='Marquer comme terminé']")
        except ElementClickInterceptedException:
            enregis = U.wait.until(EC.presence_of_element_located((By.XPATH, "//span[text()='Marquer comme terminé']")))
            U.driver.execute_script("arguments[0].click();", enregis)
            marquer = U.wait.until(EC.presence_of_element_located((By.XPATH, "//span[text()='Marquer comme terminé']")))
            U.driver.execute_script("arguments[0].click();", marquer)

        time.sleep(10)

        def enregistrer():
            cliquer("//*[text()='Enregistrer']")
            cliquer("//span[text()='Marquer comme terminé']")
            time.sleep(3)
            # Cliquer("//button[@accesskey='O']") #Oui

        U.deconnexion()
        time.sleep(7)
        U.login_1('xxx', 'xxx')
        time.sleep(7)

        U.notif(3, 4)

        for ele in ["Astr", "Phenix", "Siant", "FamilyManagement"]:
            U.val_random("//span[contains(@id, 'XXSCF" + ele + "AT::cntnrSpan')]", \
                         "//div[contains(@id, 'XXSCF" + ele + "AT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr/td[2][string-length()>1]", \
                         "//div[contains(@id, 'XXSCF" + ele + "AT::dropdownPopup::dropDownContent::db')]/descendant::tbody/tr[", \
                         "Famille " + ele + '"', True)

        dic9 = U.creation_dict(U.dict_FSC)
        print(dic9)
        dic10 = {**dic8, **dic9}
        print(dic10)

        target = U.driver.find_element_by_xpath("//select[@title='BROUILLON']")
        actions = ActionChains(U.driver)
        actions.move_to_element(target)
        actions.perform()

        U.Cliquer("//select[@title='BROUILLON']//option[@title='ACHAT']")

        time.sleep(5)
        U.Cliquer("//*[text()='OK']")
        time.sleep(5)
        enregistrer()

        lab = dic10.keys()
        lab_list = list(lab)
        #
        val = dic10.values()
        val_list = list(val)

        # fname = "C:/Users/xxx/Desktop/Achats/test1.xlsx"

        def vers_Word(num, img, doc):
            doc.add_paragraph("Approbation de l'article " + str(num))
            doc.add_picture(img, width=docx.shared.Inches(7.2), height=docx.shared.Inches(4))
            doc.save(fname_w)

        # def vers_Word_error(num, img, doc):
        #    doc.add_paragraph("Approbation de l'article " + str(num) + "incluant anomalie sous la section: ACHAT")
        #    doc.add_picture("C:/Users/xxx/Desktop/Achats/Anomalie" + result + ".png", width=docx.shared.Inches(7.2), height=docx.shared.Inches(4))
        #    doc.add_picture(img, width=docx.shared.Inches(7.2), height=docx.shared.Inches(4))
        #    doc.save(fname_w)

        def final(text):
            resultat_final = text + NIR
            date_auj = date()
            heure = heure()
            val_list.append(date_auj)
            val_list.append(heure)
            val_list.append(resultat_final)
            lab_list.append('Date')
            lab_list.append('Temps')
            lab_list.append('Resultat')
            vers_Excel(lab_list, val_list)

        i = 1
        while i <= 5:
            i += 1
            try:
                U.Cliquer("//*[contains(@id, 'UISatr:0:cil1::icon')]")
                time.sleep(3)
            except StaleElementReferenceException as e:
                U.Cliquer("//span[contains(@title,'Notifications ')]")
                time.sleep(2)
            try:
                # WebDriverWait(U.driver, 5).until(EC.presence_of_element_located((By.XPATH, "//div[@data-afrrk = '0']//td/span[contains(@title, 'secondes')]")))
                WebDriverWait(U.driver, 8).until(EC.presence_of_element_located((By.XPATH,
                                                                                 "//button[contains(@title, 'Approuver   La demande de nouvel article " + str(
                                                                                     NIR) + "')]")))
                while False:
                    break
                else:
                    # U.driver.refresh()
                    # U.Cliquer("//span[contains(@title,'Notifications ')]")
                    # U.Cliquer("//button[@id='_FOpt1:_UISatr:0:lv4:0:cb2' and text()='Approuver']")
                    U.Cliquer("//button[contains(@title, 'Approuver   La demande de nouvel article " + str(NIR) + "')]")
                    print("Test passed !")
                    U.driver.get_screenshot_as_file("C:/Users/xxx/Desktop/Achats/Approuve " + str(NIR) + ".png")
                    if not path.exists("C:/Users/xxx/Desktop/Achats/Anomalie" + result + ".png"):
                        pass
                    else:
                        doc.add_picture("C:/Users/xxx/Desktop/Achats/Anomalie" + result + ".png",
                                        width=docx.shared.Inches(7.2), height=docx.shared.Inches(4))
                    vers_Word(NIR, "C:/Users/xxx/Desktop/Achats/Approuve " + str(NIR) + ".png")
                    final('Approuve ')
                    break
            except TimeoutException as e:
                pass
            except StaleElementReferenceException as e:
                pass
        else:
            U.driver.refresh()
            U.Cliquer("//span[contains(@title,'Notifications ')]")
            U.driver.get_screenshot_as_file("C:/Users/xxx/Desktop/Achats/Non-Approuve " + str(NIR) + ".png")
            if path.exists("C:/Users/xxx/Desktop/Achats/Anomalie" + result + ".png"):
                doc.add_picture("C:/Users/xxx/Desktop/Achats/Anomalie" + result + ".png",
                                width=docx.shared.Inches(7.2), height=docx.shared.Inches(4))
            else:
                pass
            vers_Word(NIR, "C:/Users/xxx/Desktop/Achats/Non-Approuve " + str(NIR) + ".png")
            final('Non-Approuve ')
            
class fn_tnr71:
    def __init__(self, articlenum, desc, org, numero, selecter, copier, coller, msg, rapportnum, testnum):
        self.articlenum = articlenum
        self.desc = desc
        self.org = org
        self.numero = numero
        self.selecter = selecter
        self.copier = copier
        self.coller = coller
        self.msg = msg
        self.rapportnum = rapportnum
        self.testnum = testnum

    def assurer_EAN13_positif(self):
        try:
           while True:
               EAN13 = driver.find_element_by_name('_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:pt1:ap1:sspecf:1:dynReg:1:CTXRNj_ItemXxscfOrangeIcPrivateVOXXSCF_EAN_GENERATION_AG:0:XXSCFEANUsageAT')
               if EAN13 == 'N':
                   EAN13.clear()
                   EAN13.send_keys('O')
                   break
        except Exception:
           print("EAN13 est deja OUI")

    def obtenir_informations(self):  # fonction pour recuperer les valeurs et les presentes
        global xpath_values1
        nom_keys = ['CM', 'MC', 'COU', 'MCode', 'NA', 'LM', 'TR', 'SAV', 'TDS', 'CPS', 'CDR', 'UEAN13', 'CARTON', 'PALETTE', 'EAN13A', 'Carton',
                    'Palette', 'ment_fourni', 'ref_fourni', 'fourni', 'PMS', 'UAeD', 'CapMem', 'NivTra', 'TraAtt', 'INI', 'DebVali']
        # Pour trouver les valeurs des champs
        xpath_values1 = ['BrandAT_Display', 'CommercialModelAT', 'ColorAT', 'ModelCodeAT', 'MobileNaCodeAT', 'LabelModelAT', 'ReferenceTypeAT',
                        'SavEligibilityAT', 'InvDepreciationRateAT', 'CPSAT', 'CDRAT', 'EANUsageAT', 'GenerCartonEANAT', 'GenerPalletEANAT',
                        'EANItemAT', 'EanCartonAT', 'EanPalletAT', 'SupplierRefBoolAT', 'SupplierRefAT', 'SupplierAT_Display', 'PmsEligibilityAT',
                        'UsingArticleInDemoAT', 'MemoryCapacityAT', 'TraceabilityAT', 'ExpTraceabilityCodeAT', 'StartDateAT', 'ValidFromDateAT']
        # Pour trouver les Labels
        xpath_values2 = ['Marque', 'Modèle C', 'Couleur', 'code', 'NA', 'Libellé', 'Type R', 'SAV', 'dépréciation', 'CPS', 'CDR', 'EAN13',
                        'CARTON', 'PALETTE', 'N 13 Article', 'N Carton', 'N Palette', 'Référencement', 'Référence f', 'Fournisseur', 'PMS',
                        'demo', 'Mémoire', 'Niveau', 'attendue', 'initialisation', 'validité']

        global list_label, list_valeurs, dictio
        list_valeurs = []
        list_label = []
        dictio = dict(zip(nom_keys, xpath_values1))  # Creer une dictionaire a partir des deux listes
        dictio1 = dict(zip(nom_keys, xpath_values2))
        #l.log("INFO", "Le dictionnaire pour les valeurs est: " + str(dictio), l.log_file)
        #l.log("INFO", "Le dictionnaire pour les labels est: " + str(dictio1), l.log_file)
        print("Le dictionnaire pour les valeurs est: ", str(dictio))
        print("Le dictionnaire pour les labels est: ", str(dictio1))

        for key, value in dictio.items():
            key1 = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, '" + value + "')]"))).get_attribute('value')
            #key1 = wait.until(EC.presence_of_element_located((By.XPATH, "//span[contains(@id, '"+value+"')]"))).get_attribute('value')
            list_valeurs.append(key1)  # Obtenir les valeurs des differents champs
        print("La liste des valeurs est: ", list_valeurs, "\n")
        #return list_valeurs

        for key, value in dictio1.items():
            key2 = wait.until(EC.presence_of_element_located((By.XPATH, "//label[contains(text(), '" + value + "') and contains(@class, 'label-text')]"))).text
            list_label.append(key2)
        list_label.pop(19)
        list_label.insert(19, 'Fournisseur')
        print("La liste des labels est: ", list_label, "\n")

        dictio2 = dict(zip(list_label, list_valeurs))
        print("Le dictionnaire avec des valeurs pour l'article 227481 est: ", dictio2, "\n")

        def display_value(titre, startI, endI):
            print("Les valeurs pour " + titre + ": ")
            for i in range(startI, endI):
                print(list_label[i], ": ", list_valeurs[i])
            print()

        display_value('Description Commerciale', 0, 6)
        display_value('Informations générales', 6, 11)
        display_value('Génération EAN', 11, 14)
        display_value('EAN', 14, 17)
        display_value('Référence fournisseur', 17, 20)
        display_value('Caractéristiques Produit', 20, 23)
        display_value('Traçabilité', 23, 25)
        display_value('Date', 25, 27)

    def copiercoller(xpath, numero, selecter, copier, coller):
       global EAN13Article_original
       EAN13Article = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, xpath)))
       EAN13Article_original = EAN13Article.get_attribute('value')
       print("Valeur EAN13 Article de l'article " + numero + " est: ", EAN13Article_original)
       EAN13Article.send_keys(Keys.CONTROL, "'" + selecter + "'")
       EAN13Article.send_keys(Keys.CONTROL, "'" + copier + "'")
       try:
           if coller != '':
               EAN13Article.clear()
               EAN13Article.send_keys(Keys.CONTROL, "'"+coller+"'")
       finally:
           if copier == 'c':
               print("L'EAN13 a été copiée")
           elif coller == 'v':
               print("L'EAN13 est maintenant collée dans l'article 226352")

    def rapport_test(msg, numero):
       pass_message = "Test " + str(numero) + " réussi: Echec du sauvegarde. " + msg
       fail_message = "Test " + str(numero) + " échoué: Aucun message d'alerte! "
       Chemin = "C:/Users/" + os.getlogin() + "/Desktop/71/"
       try:
           ok = WebDriverWait(driver, 30).until(EC.element_to_be_clickable((By.XPATH, "//button[@_afrpdo = 'cancel' and text()='OK']")))
           if ok.is_displayed():
               time.sleep(2)
               driver.get_screenshot_as_file(Chemin + str(numero) + ".png")
               l.logger.info(pass_message)
               ok.click()
               print()
       except Exception as e:
           driver.get_screenshot_as_file(Chemin + str(numero) + ".png")
           l.logger.error(fail_message)
           print()

    def testsave(testnum):
       try:
           ok = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//button[@_afrpdo = 'cancel' and text()='OK']")))
           def report(msg1, msg2):
               time.sleep(2)
               driver.save_screenshot(Directory + "/" + str(testnum) + "save.png")
               fonctionnel.vers_Word(msg1, Directory + "/" + str(testnum) + "save.png")
               fn_tnr71.rapport_test(msg2, testnum)
           def report1(msg1, msg2):
               fonctionnel.vers_Word(msg1, Directory + "/" + str(testnum) + "notsave.png")
               fn_tnr71.rapport_test(msg2, testnum)
           if ok.is_displayed() and testnum == 1:   # Premier test avec seulement EAN 13 Article modifiée
               report("Sauvegarde 'ECHEC' avec l'EAN13 de 227481 insérée dans celle de 226352. Test passed!",
                      "\nEAN 13 fournisseur déjà présent dans la base SISTER.\n"
                      "L'EAN 13 a été remplacée par celle de l'article 227481.\n")
               
           elif ok.is_displayed() and testnum == 2:
               report("Sauvegarde 'ECHEC' avec d'autres valeurs modifiées. Test passed!",
                      "\nEAN 13 fournisseur déjà présent dans la base SISTER. \n"
                      "Les valeurs ont été modifiées pour qu'elles correspondent à celle de l'article 227481.\n")
               
           elif not(ok.is_displayed()) and testnum == 1:
               driver.save_screenshot(Directory + "/" + str(testnum) + "notsave.png")
               report1("Sauvegarde 'PASS' avec l'EAN13 de l'article 227481 insérée dans celle de 226352. Test failed!",
                       "\nEAN 13 fournisseur déjà présent dans la base SISTER.\n")
           elif not(ok.is_displayed()) and testnum == 2:
               driver.save_screenshot(Directory + "/" + str(testnum) + "notsave.png")
               report1("Sauvegarde 'PASS' avec d'autres valeurs modifiées. Test failed!",
                       "\nEAN 13 fournisseur déjà présent dans la base SISTER. \n"
                       "Les valeurs ont été modifiées pour qu'elles correspondent à celle de l'article 227481.\n")
       finally:
           print("Rapport disponible")

    def supprimer_remplacer_info(self):
       xpath_v2 = []
       for xp in xpath_values1:
           xpath_v2.append("//input[contains(@name, '" + xp + "')]")

       for i in range(len(list_valeurs)):
           scroll(xpath_v2[i])
           champs = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, xpath_v2[i])))
           champs1 = champs.get_attribute('value')
           if champs1 == list_valeurs[i]:
               print("Le champs pour ", list_label[i], " est identique pour les 2 articles dont la valeur est: ", champs1, "\n")
               time.sleep(1)
           elif champs1 != list_valeurs[i]:
               print("Le champs pour ", list_label[i], " N'EST PAS identique pour les 2 articles dont la valeur est: ", champs1)
               champs.clear()
               champs.send_keys(list_valeurs[i])

    def dernier_test(self):
        try:
           ok = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//button[@_afrpdo = 'cancel' and text()='OK']")))
           if ok.is_displayed():
               driver.save_screenshot(Directory + "/final.png")
               fonctionnel.vers_Word("Confirmation de sauvegarde échec avec l'EAN 13 originale pour l'article 226352", Directory + "/final.png")
               l.logger.info("Test 3 failed: Sauvegarde échoué.\n"
                             "Les valeurs ont été modifiées pour qu'elles correspondent à celle de l'article 227481 et l'EAN13 \n"
                             "Article(4718487703333) pour l'article 226352 a été retenue avant de sauvegarder. A noter que Eligibilité PMS \n"
                             "et Tracabilite Attendue n'ont pas été affectés d'autre valeurs dans leurs champs respectifs parceque le message \n"
                             "d'alerte afficheront que ces champs sont obligatoires meme s'ils ont déja été remplacés par d'autres valeurs. \n"
                             "Sauvegarde abouti !\n ************************************************************")
        except TimeoutException as e:
               driver.save_screenshot(Directory + "/final.png")
               time.sleep(5)
               fonctionnel.vers_Word("Confirmation de sauvegarde réussi avec l'EAN 13 originale pour l'article 226352", Directory + "/final.png")
               l.logger.info("Test 3 réussi: Sauvegarde effectué.\n"
                     "Les valeurs ont été modifiées pour qu'elles correspondent à celle de l'article 227481 et l'EAN13 \n"
                     "Article(4718487703333) pour l'article 226352 a été retenue avant de sauvegarder. A noter que Eligibilité PMS \n"
                     "et Tracabilite Attendue n'ont pas été affectés d'autre valeurs dans leurs champs respectifs parceque le message \n"
                     "d'alerte affichera que ces champs sont obligatoires meme s'ils ont déja été remplacés par d'autres valeurs. \n"
                     "Sauvegarde abouti !\n \n **************************************************************************************************\n")
        finally:
           print("Rapport final disponible !")

class fn_tnr72:
    def __init__(self, numero, chemin, min_num, max_num, no_chiffre):
        self.numero = numero
        self.chemin = chemin
        self.min_num = min_num
        self.max_num = max_num
        self.no_chiffre = no_chiffre

    def choisir_FEH_IO_MASTER(numero):
        cliquer("//div[@title='Tâches']")
        cliquer("//*[text()='Gérer les articles']")
        # Rechercher le numero de l'article
        Article = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, '_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:pt1:ItemC1:0:simplePanel1:region2:0:efqrp:value00')]")))
        Article.send_keys(numero)
        press_enter()

        wait.until(EC.presence_of_all_elements_located((By.XPATH, "//table[contains(@summary, 'Résultats de la recherche d')]")))
        driver.get_screenshot_as_file("C:/Users/xxx/Desktop/tnr727576/Org_non-affiche.png")
        wait.until(EC.presence_of_all_elements_located((By.XPATH, "//table[contains(@summary, 'Résultats de la recherche d')]//tbody/tr/td/span/a[text()='" + str(numero) + "']")))
        cliquer("//tr[@_afrrk = '0']//a[text() = '" + str(numero) + "']")
        time.sleep(5)

    def modifier_utilisationEAN13_par_oui(self):
        Util_EAN13 = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, 'ItemXxscfOrangeIcPrivateVOXXSCF_EAN_GENERATION_AG:0:XXSCFEANUsageAT')]")))
        Util_EAN13_valeur = Util_EAN13.get_attribute('value')
        if Util_EAN13_valeur == 'O':
            pass
        else:
            Util_EAN13.clear()
            Util_EAN13.send_keys('O')

    def nombre_chiffre_test(chemin, min_num, max_num, no_chiffre):
        EAN13_article = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, 'ItemXxscfOrangeIcPrivateVOXXSCF_EAN_AG:0:XXSCFEANItemAT')]")))
        EAN13_article.clear()
        chiffre = random.randint(min_num, max_num)
        EAN13_article.send_keys(chiffre)
        time.sleep(3)
        driver.get_screenshot_as_file(chemin + "/EAN13_A_" + str(no_chiffre) + " chiffres.png")
        fonctionnel.vers_Word("CREATION ARTICLE EAN13 AVEC " + str(no_chiffre) + " CHIFFRES", chemin + "/EAN13_A_" + str(no_chiffre) + " chiffres.png")

        scroll("//*[text() = 'Sauvegarder']")
        cliquer("//*[text() = 'Sauvegarder']")
        time.sleep(5)
        try:
            ok = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//button[@ID = '_FOd1::msgDlg::cancel' and text() = 'OK']")))
            # erreur = U.WebDriverWait(U.driver, 5).until(EC.presence_of_element_located((By.XPATH, "//td[text()='Erreur']")))
            if ok.is_displayed():
                driver.get_screenshot_as_file(chemin + "/ECHEC_EAN13_A_" + str(no_chiffre) + " chiffres.png")
                fonctionnel.vers_Word("Test reussi: REJET de l'article EAN13 avec " + str(no_chiffre) + " chiffres", chemin + "/ECHEC_EAN13_A_" + str(no_chiffre) + " chiffres.png")

        except Exception as e:
            driver.get_screenshot_as_file(chemin + "/ECHEC_erreur_EAN13_A_" + str(no_chiffre) + " chiffres.png")
            fonctionnel.vers_Word("Test reussi: REJET de l'article EAN13 avec " + str(no_chiffre) + " chiffres", chemin + "/ECHEC_erreur_EAN13_A_" + str(no_chiffre) + " chiffres.png")

    def chiffre_a_13_errones(chemin):
        EAN13_article = driver.find_element_by_xpath("//input[contains(@name, 'ItemXxscfOrangeIcPrivateVOXXSCF_EAN_AG:0:XXSCFEANItemAT')]")
        EAN13_article.clear()
        no = '356129' + str(random.randint(1000_000, 10_000_000))
        EAN13_article.send_keys(no)
        time.sleep(3)
        driver.get_screenshot_as_file(chemin + "/EAN13_" + str(13) + "chiffres_commencant_par_356129.png")
        fonctionnel.vers_Word("CREATION ARTICLE EAN13 AVEC 13 CHIFFRES COMMENCANT PAR '356129'", chemin + "/EAN13_" + str(13) + "chiffres_commencant_par_356129.png")

        scroll("//*[text() = 'Sauvegarder']")
        time.sleep(2)
        cliquer("//*[text() = 'Sauvegarder']")
        time.sleep(5)
        try:
            ok = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, "//button[@ID = '_FOd1::msgDlg::cancel' and text() = 'OK']")))
            if ok.is_displayed():
                driver.get_screenshot_as_file(chemin + "/ECHEC_EAN13_" + str(no) + ".png")
                fonctionnel.vers_Word("Test reussi: REJET de l'article EAN13 " + str(no), chemin + "/ECHEC_EAN13_" + str(no) + ".png")
        except Exception as e:
            driver.get_screenshot_as_file(chemin + "/ECHEC_erreur_EAN13_" + str(no) + ".png")
            fonctionnel.vers_Word("Test reussi: REJET de l'article EAN13 " + str(no), chemin + "/ECHEC_erreur_EAN13_" + str(no) + ".png")

class fn_tnr7_256:
    def __init__(self, navigateur_ele, option, Date, optionD, optionP, optionTOF, desc, row):
        self.navigateur_ele = navigateur_ele
        self.option = option
        self.Date = Date
        self.optionD = optionD
        self.optionP = optionP
        self.optionTOF = optionTOF
        self.desc = desc
        self.row = row

    def naviguer_vers_Accueil(navigateur_ele, option):
        try:
            cliquer("//a[@title = 'Navigateur']")
        except TimeoutException as e:
            driver.refresh()
            cliquer("//a[@title = 'Navigateur']")
        cliquer(navigateur_ele)
        time.sleep(3)
        cliquer(option)

    def alimentation_des_champs(Date, optionD, optionP, optionTOF, desc):
        # Cliquer("//*[contains(@title, 'Champ flexible extensible d') and @class = 'x2h']")
        date = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@name = '_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value00']")))
        date.clear()
        date.send_keys(Date)
        DATE = driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:operator0::content')
        drp = Select(DATE)
        drp.select_by_visible_text(optionD)
        # PRODUIT
        Produit = driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value20::content')
        drp = Select(Produit)
        drp.select_by_visible_text(optionP)
        # Cliquer("//select[@name = '_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value40']/descendant::option[contains(@title, " + option + ")]")
        # Cliquer("//span[@id = '_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value40']")
        # Cliquer("//select[@id = '_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value40::content']")
        try:
            time.sleep(7)
            TOF = driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value40::content')
        finally:
            TOF.send_keys(optionTOF)

        description = driver.find_element_by_xpath("//input[@name = '_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:qryId1:value50']")
        description.clear()
        description.send_keys(desc)
        cliquer("//button[@accesskey = 'e' and text() = 'chercher']")
        try:
            WebDriverWait(driver, 50).until(EC.visibility_of_any_elements_located((By.XPATH, "//table[@summary = 'Résultats de la recherche']")))
        finally:
            time.sleep(2)
            driver.get_screenshot_as_file("C:/Users/xxx/Desktop/tnr727576/tnr76/Article_modifie_check.png")
        time.sleep(3)

    def cocher_tous_les_attributs(self):
        checkbox = driver.find_element_by_xpath("//input[@name = '_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:sbc1']")
        actions = ActionChains(driver)
        actions.move_to_element(checkbox)
        actions.click(checkbox).perform()

        cliquer("//a[@title = 'Recherche : Sélectionner un attribut']")
        cliquer("//a[text() = 'Tous les attributs']")
        try:
            WebDriverWait(driver, 25).until(EC.presence_of_element_located((By.XPATH, "//input[@value= 'Tous les attributs']")))
        except TimeoutException as e:
            WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//input[@value= 'Tous les attributs']")))
        finally:
            print("Sélectionner 'Tous les attributs': Activé")

    def telecharger_les_donnees(self):
        user = os.getlogin()
        fname = "C:/Users/" + user + "/Downloads/auditreport.csv"
        dir_exists = os.path.isdir(fname)
        if dir_exists:
            print("Le fichier csv est présent dans le bureau.")
        cliquer("//a[text() = 'Actions']")
        cliquer("//*[text() = 'Exporter au format CSV']")

        secondes = 0
        while not os.path.exists(fname) and secondes <= 50:
            print("Le fichier n'est pas encore téléchargé. Veuillez patienter !")
            time.sleep(2)
            secondes += 1
            if os.path.exists(fname):
                print('Fichier csv téléchargé !')
                break
            else:
                pass

    def scrape_toExcel(row, nb_colonnes, liste_titres, chemin):
        nb_lignes = len(driver.find_elements_by_xpath("//table[@summary = 'Résultats de la recherche' and @_startrow = '" + str(row) + "']//tr"))
        for i in range(1, nb_lignes + 1):  #
            liste_valeurs = []
            for j in range(1, nb_colonnes + 1):
                # U.scroll("//table[@summary = 'Résultats de la recherche']//tbody/tr["+str(i)+"]/td["+str(j)+"]")
                scroll("//table[@summary = 'Résultats de la recherche' and @_startrow = '" + str(row) + "']//tbody/tr[" + str(i) + "]")
                valeur = wait.until(EC.presence_of_element_located((By.XPATH,
                        "//table[@summary='Résultats de la recherche' and @_startrow='" + str(row) + "']//tbody/tr[" + str(i) + "]/td[" + str(j) + "]"))).text
                liste_valeurs.append(valeur)
            # print(liste_valeurs)
            fonctionnel.vers_Excel(chemin, liste_titres, liste_valeurs)

    # ['_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c2','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c1','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c8','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c7','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c9','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c6','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c5','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c4']
    def extraction_des_donnees_auto(self):
        Jour = datetime.today().strftime('%Y-%m-%d')
        Heure = time.strftime("%H.%M.%S", time.localtime())

        nb_colonnes = len(driver.find_elements_by_xpath("//table[@summary = 'Cette table contient des en-têtes de colonne correspondant à la table de corps de données ci-dessous']//tbody/tr/th[@align = 'left']"))
        liste_titres = []

        k = 1
        while k <= nb_colonnes:
            titre = wait.until(EC.presence_of_element_located((By.XPATH,
                    "//table[contains(@summary, ' correspondant à la table de corps de données ci-dessous')]//tbody/tr/th[@align = 'left'][" + str(k) + "]")))
            liste_titres.append(titre.text)
            k += 1
        # ['_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c2','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c1','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c8','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c7','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c9','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c6','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c5','_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:AT1:_ATp:ATt1:c4']
        for row in range(0, 100000, 25):
            try:
                WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.XPATH, "//table[@summary = 'Résultats de la recherche' and @_startrow = '" + str(row) + "']//tbody/tr[20]/td[1]")))
                # U.WebDriverWait(U.driver, 5).until(EC.presence_of_element_located((By.XPATH, "//table[@summary = 'Résultats de la recherche' and @_startrow = '"+str(row)+"']//tbody/tr[25]/td[1]")))
                while True:
                    scroll("//table[@summary = 'Résultats de la recherche' and @_startrow = '" + str(row) + "']//tbody/tr[25]/td[1]")
                    fn_tnr7_256.scrape_toExcel(row, nb_colonnes, liste_titres, Directory + "/tnr76_" + Jour + "_" + Heure + ".csv")
                    print(row)
                    break
            except TimeoutException as e:
                nb_lignes1 = len(driver.find_elements_by_xpath("//table[@summary = 'Résultats de la recherche' and @_startrow = '" + str(row) + "']//tr"))
                print(nb_lignes1)
                fn_tnr7_256.scrape_toExcel(row, nb_colonnes, liste_titres, Directory + "/tnr76_" + Jour + "_" + Heure + ".csv")
                driver.get_screenshot_as_file(Directory + "/end.png")
                fonctionnel.vers_Word("Fin de l'extraction des données! Au bout de la liste déroulante.", Directory + "/end.png")
                print("Fin !")
                break

class fn_tnr69_8_67:
    def __init__(self, modifier_les_options_de_lot_article, lot, utilisateur, num_tnr, afrrk, csv_nom, mappe_importation, num_article, page_gestion_article_disponible):
        self.modifier_les_options_de_lot_article = modifier_les_options_de_lot_article
        self.num_tnr = num_tnr
        self.afrrk = afrrk
        self.csv_nom = csv_nom
        self.mappe_importation = mappe_importation
        self.utilisateur = utilisateur
        self.lot = lot
        self.num_article = num_article
        self.page_gestion_article_disponible = page_gestion_article_disponible

    def gestion_des_articles(self, lot, num_tnr):
        cliquer("//a[@title='Navigateur']")
        cliquer("//div[@title = 'Gestion des produits']")
        cliquer("//a[@title='Gestion des informations produit']")
        cliquer("//div[@title='Tâches']")
        if lot == False:
            cliquer("//*[text()='Gérer les articles']")
            cliquer("//button[text()='Rechercher']")
            #if num_art != '':
            #    Article = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@name, 'ItemC1:0:simplePanel1:region2:0:efqrp:value00')]")))
            #    Article.send_keys(num_art)
        elif lot == True:
            cliquer("//*[contains(text(), 'Gérer les lots d')]")
            Nom = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'pt1:AP1:r1:0:q1:value10::content')]")))
            if num_tnr == '69' or num_tnr == '86':
                Nom.send_keys('creation')
            elif num_tnr == '87':
                Nom.send_keys('modif_taux_depre')
            Affecter_a = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'AP1:r1:0:q1:value20::content')]")))
            Affecter_a.clear()
            cliquer("//button[contains(@id, '0:pt1:AP1:r1:0:q1::search')]")  # Rechercher

    def reaffectation_lot_article_au_bon_utilisateur(self, utilisateur, num_tnr):
        try:
            if num_tnr == '69' or num_tnr == '86':
                WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//a[text() = 'creation_article_mobile']")))
                while True:
                    print("Resultat lot d'article obtenu pour le tnr", num_tnr)
                    break
            elif num_tnr == '87':
                WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//a[text() = 'modif_taux_depre']")))
                while True:
                    print("Resultat lot d'article obtenu pour le tnr", num_tnr)
                    break

        except TimeoutException as e:
            print("Resultat lot d'article null. Nouvelle recherche avec la date parametrer a: Egal ou antérieur à")
            cliquer("//select[@title = 'Egal ou postérieur à']")
            cliquer("//select[@title = 'Egal ou postérieur à']/option[@title = 'Egal ou antérieur à']")
            cliquer("//button[contains(@id, '0:pt1:AP1:r1:0:q1::search')]")  # Rechercher
            time.sleep(5)
            try:
                if num_tnr == '69' or num_tnr == '86':
                    WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//a[text() = 'creation_article_mobile']")))
                    while True:
                        print("Resultat lot d'article obtenu pour le tnr ", num_tnr," apres une nouvelle recherche !")
                        break
                    else:
                        pass
                elif num_tnr == '87':
                    WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//a[text() = 'modif_taux_depre']")))
                    while True:
                        print("Resultat lot d'article obtenu pour le tnr ", num_tnr," apres une nouvelle recherche !")
                        break
            finally:
                print("Recherche lot d'article complet")
        '''    
        try:
            WebDriverWait(driver, 17).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'CDG']")))
            while True:
                print("Resultat obtenu. Le lot d'article a ete affecter a CDG.")
                break
        except TimeoutException as e:
            WebDriverWait(driver, 8).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'RGA']")))
            while True:
                print("Resultat obtenu. Le lot d'article a ete affecter a RGA.")
        except TimeoutException as e:
            print("Aucun resultat obtenu pour le lot d'article! ")
            cliquer("//select[@title = 'Egal ou postérieur à']/option[@title = 'Egal ou antérieur à']")
            cliquer("//select[@title = 'Egal ou postérieur à']")
            cliquer("//button[contains(@id, '0:pt1:AP1:r1:0:q1::search')]")  # Rechercher
        '''
        #print("Attente...")
        #try:
        #    #WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//div[contains(@id, 'pt1:AP1:r1:0:SearchResultsTable:_ATp:blt::db') and text() = 'Aucun résultat trouvé.']")))
        #    WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'CDG']"))) or \
        #    WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'RGA']")))
        #    while True:
        #        print("Resultat obtenu pour le lot d'article")
        #        break
        #except TimeoutException as e:
        #    cliquer("//a[@title = 'Réduire Recherche avancée']")
        #    cliquer("//select[@title = 'Egal ou postérieur à']")
        #    cliquer("//select[@title = 'Egal ou postérieur à']/option[@title = 'Egal ou antérieur à']")
        #    cliquer("//button[contains(@id, '0:pt1:AP1:r1:0:q1::search')]")

            #no_result = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//*[text() = 'Aucun résultat trouvé.']")))
        # C'est soit 1 ou 2 qui est éligible
        def reaffecter():
            try:
                cliquer("//a[text() = 'Actions']")
            except StaleElementReferenceException as e:
                cliquer("//a[text() = 'Actions']")
            cliquer("//td[text() = 'Réaffecter']")
            Affecter = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, ':SearchResultsTable:assigneeIdDispId::content')]")))
            time.sleep(3)
            Affecter.clear()
            Affecter.send_keys(utilisateur)
            time.sleep(2)
            cliquer("//*[text() = 'auvegarder et fermer']")
            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/Affecter.png")

        def effectuer_la_bonne_reaffectation():
            #cliquer("//a[@title = 'Réduire Recherche avancée']")
            #cliquer("//select[@title = 'Egal ou postérieur à']")  # --> (1) <--
            #cliquer("//button[text() = 'Rechercher']")
            cliquer("//table[contains(@summary, 'Lots d')]")
            time.sleep(3)
            if num_tnr == '69' or num_tnr == '86':
                cliquer("//span[text() = '22007']")
            elif num_tnr == '87':
                cliquer("//span[text() = '7004']")
            time.sleep(5)
            reaffecter() # *
            #except TimeoutException as e:
            #    cliquer("//a[@title = 'Développer Recherche avancée']")
            #    cliquer("//select[@title = 'Egal ou postérieur à']")
            #    #cliquer("//option[@value = 'ONORBEFORE' and @title = 'Egal ou antérieur à']")  # --> (2) <--
            #    cliquer("//button[text() = 'Rechercher']")
            #    time.sleep(5)
            #    cliquer("//span[text() = '22007']")
            #    reaffecter() # *

        try:
            statut_CDG = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'CDG']")))
            while True:
                print(statut_CDG.text, "TROUVE !")
                if num_tnr == '87':
                    print("L'affectation est deja bonne pour CDG. Pas de nouvelle réaffectation!")
                    scroll("//span[text() = 'CDG']")
                    driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                    fonctionnel.vers_Word("L'affectation est deja la bonne pour " + str(utilisateur),
                                          "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                    #effectuer_la_bonne_reaffectation()
                    break
                elif num_tnr == '69' or num_tnr == '86':
                    print("L'affectation n'est pas bonne pour RGA")
                    driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                    fonctionnel.vers_Word("L'affectation n'est pas la bonne pour " + str(utilisateur),
                                          "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                    effectuer_la_bonne_reaffectation()
                    fonctionnel.vers_Word("Nouvelle réaffectation du lot d'article pour " + str(utilisateur),
                                          "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Affecter.png")
                    break
        except TimeoutException as e:
            print("CDG n'etait pas affecte auparavant")
            try:
                statut_RGA = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'RGA']")))
                while True:
                    print(statut_RGA.text, "TROUVE !")
                    if num_tnr == '87':
                        print("L'affectation n'est pas bonne pour le lot d'article modif_tax_depre")
                        driver.get_screenshot_as_file("C:/Users/"+os.getlogin()+"/Desktop/TNR-"+str(num_tnr)+"/C_A_M_G.png")
                        fonctionnel.vers_Word("L'affectation n'est pas la bonne pour "+str(utilisateur),
                                              "C:/Users/"+os.getlogin()+"/Desktop/TNR-"+str(num_art)+"/C_A_M_G.png")
                        effectuer_la_bonne_reaffectation()
                        fonctionnel.vers_Word("Nouvelle réaffectation du lot d'article pour CDG",
                                              "C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/Affecter.png")
                        break
                    elif num_tnr == '69' or '86':
                        print("L'affectation est deja bonne pour RGA")
                        driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                        fonctionnel.vers_Word("L'affectation est deja la bonne pour " + str(utilisateur),
                                              "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                        #effectuer_la_bonne_reaffectation()
                        break
            except TimeoutException as e:
                print("CDG et RGA n'etaient pas affectes auparavant")
                try:
                    statut = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//span[contains(text(), '@orange.com')]")))
                    while True:
                        print(statut.text, "TROUVE !")
                        if num_tnr == '87':
                            print("L'affectation n'est pas bonne pour le lot d'article modif_tax_depre")
                            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                            fonctionnel.vers_Word("L'affectation n'est pas la bonne pour " + str(utilisateur),
                                                  "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                            effectuer_la_bonne_reaffectation()
                            fonctionnel.vers_Word("Nouvelle réaffectation du lot d'article pour modif_tax_depre",
                                                  "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Affecter.png")
                            break
                        #elif num_tnr == '69' or '86':
                        #    print("L'affectation est deja bonne pour creation_article_mobile")
                        #    driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                        #    fonctionnel.vers_Word("L'affectation est deja la bonne pour " + str(utilisateur),
                        #                          "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/C_A_M_G.png")
                        #    # effectuer_la_bonne_reaffectation()
                        #    break
                finally:
                    print("Fin des affectations de lot d'article respectif")

    def clique_Action(self, num_tnr): # 69 86 87
        #global nu
        if num_tnr == '69' or num_tnr == '86':
            try:
                cliquer("//table[contains(@summary, 'Lots d')]//tr[@_afrrk = '0']")
            except (StaleElementReferenceException, TimeoutException) as e:
                cliquer1("//table[contains(@summary, 'Lots d')]//tr[@_afrrk = '1']")
        elif num_tnr == '87':
            try:
                cliquer("//div[contains(@id, 'SearchResultsTable:_ATp:blt::db')]")
            except StaleElementReferenceException as e:
                cliquer1("//span[text() = '7004']")

        #numero = driver.find_element_by_xpath("//table[contains(@summary, 'Lots d')]//td[1]/span[@class = 'x2o2']")
        #num = driver.find_element_by_xpath()
        #nu = numero.text
        #print(nu)

        time.sleep(3)
        try:
            cliquer("//div[@aria-label = 'Actions']")
        except (TimeoutException, StaleElementReferenceException) as e:
            cliquer1("//div[@aria-label = 'Actions']")

    def purger_toutes_les_lignes(self, modifier_les_options_de_lot_article, num_tnr):
        def purgation_des_lignes():
            cliquer("//td[text() = 'Purger']")
            scroll("//td[text() = 'Toutes les lignes']")
            cliquer("//td[text() = 'Toutes les lignes']")
            time.sleep(5)
            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/purger.png")
            fonctionnel.vers_Word("Confirmation de la suppression...",
                                  "C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/purger.png")
            cliquer("//button[@accesskey = 'O']")
            time.sleep(5)
            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/purger_confirmer.png")
            fonctionnel.vers_Word("Confirmation de la suppression : Succes",
                                  "C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/purger_confirmer.png")
            cliquer("//button[@accesskey = 'K']")
            try:
                WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'Error']")))
                while True:
                    print("Erreur dans la purgation de toutes les lignes")
                    driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/Erreur_purger_lignes.png")
                    fonctionnel.vers_Word("Interdiction de purger les enregistrement du lot d'article 7004",
                                          "C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/Erreur_purger_lignes.png")
                    cliquer1("//button[@accesskey = 'K']")
                    break
            except TimeoutException as e:
                print("La purgation des lignes a ete un succes")

        if modifier_les_options_de_lot_article == False:
            purgation_des_lignes()

        elif modifier_les_options_de_lot_article == True:
            print("Il faut modifier les options de lot d'article")
            purgation_des_lignes()
            time.sleep(5)
            fn_tnr69_8_67.clique_Action(self, '86')
            scroll("//td[contains(text(), 'Modifier les options du lot d')]")
            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/Modifier_options_lot_article.png")
            fonctionnel.vers_Word("Modifier les options du lot d'article",
                                  "C:/Users/" + os.getlogin() + "/Desktop/TNR-"+str(num_tnr)+"/Modifier_options_lot_article.png")
            cliquer("//td[contains(text(), 'Modifier les options du lot d')]")

    def parametrage(self, tnr_69 = True):
        try:
            cliquer("//select[@title = 'Au chargement des données']") #1
            Planifier = wait.until(EC.element_to_be_clickable((By.XPATH, "//select[@title = 'Au chargement des données']")))
            print("Planifier est correctement: Au chargement des données \n")
        except Exception as e:
            cliquer("//select[contains(@title, 'Au chargement des données')]/option[value = '0']")
            print("Planifier est correctement: Au chargement des données \n")

        # Qualite des donnees
        # check = U.driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:r1:0:sbc5::content').is_selected()
        check = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'r1:0:sbc5::content')]"))).is_selected()
        if check == True:
            print('Confirmer les articles sans correspondance comme des nouveaux articles: COCHER OUI \n')
        elif check == False:
            print('Confirmer les articles sans correspondance comme des nouveaux articles: COCHER NON ! \n')
            cliquer("//label[@for = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:r1:0:sbc5::content']") #2

        # Ordre de modification
        if tnr_69 == False: # tnr86
            cliquer("//select[contains(@id, 'ChangeOrderOption::content')]")
            cliquer("//select[contains(@id, 'ChangeOrderOption::content')]/option[@value = '1']") #3
            cliquer("//label[text() = 'Ordre de modification']")
            ajouter_tous_les_articles = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, '0:r1:0:sbc2::content')]"))).is_selected()
            if ajouter_tous_les_articles == True: # Decocher
                print('Ajouter tous les articles: COCHER OUI \n')
                cliquer("//label[@for = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:r1:0:sbc2::content']")
                print('Ajouter tous les articles: COCHER NON maintenant ! \n')
            elif ajouter_tous_les_articles == False:
                print('Ajouter tous les articles: COCHER est deja NON !')

            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-87/Parametrage.png")
            cliquer("//span[text() = 'auvegarder et fermer']")
            fonctionnel.vers_Word("Conformite du parametrage", "C:/Users/" + os.getlogin() + "/Desktop/TNR-87/Parametrage.png")
            time.sleep(10)
        else: # tnr69
        # Demande de nouvel article
            cliquer("//select[contains(@id, 'NirOption::content')]")
            cliquer("//select[contains(@id, 'NirOption::content')]/option[@value = '1']") #3
            cliquer("//label[text() = 'Demande de nouvel article']")
            check1 = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, ':0:r1:0:sbc1::content')]"))).is_selected()
            if check1 == True:
                print('Ajouter tous les articles: COCHER OUI \n')
            elif check1 == False:
                print('Ajouter tous les articles: COCHER NON !')
                cliquer("//label[@for = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:r1:0:sbc1::content']") #4
                print('Ajouter tous les articles: COCHER OUI suite action clique ! \n')
            # Ordre de modification
            try:
                check2 = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, ':0:r1:0:sbc7::content')]"))).is_selected()
                if check2 == True:
                    print('Fractionner automatiquement: COCHER OUI \n')
                elif check2 == False:
                    print('Fractionner automatiquement: COCHER NON !')
                    cliquer("//label[@for = '_FOpt1:_FOr1:0:_FONSr2:0:MAt3:0:r1:0:sbc7::content']")
                    print('Fractionner automatiquement: COCHER OUI suite action clique ! \n')
            except TimeoutException as e:
                print("Ordre de modification: TimeoutException a eu lieu")
                check3 = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//span[contains(@id, ':0:r1:0:sbc7::content')]/span[@class = 'x10d']")))
                if check3 == True:
                    print("Impossible de cocher ou decocher Fractionner automatiquement")
            driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-69/Parametrage.png")
            cliquer("//span[text() = 'auvegarder et fermer']")
            fonctionnel.vers_Word("Conformite du parametrage", "C:/Users/"+os.getlogin()+"/Desktop/TNR-69/Parametrage.png")
            time.sleep(10)

    def chargement_fichier(self, num_tnr, csv_nom, mappe_importation):
        try:
            scroll("//td[text() = 'Ajouter des articles au lot']")
            cliquer("//td[text() = 'Ajouter des articles au lot']")
            cliquer("//span[contains(@id, 'SearchResultsTable:r4:0:mapNameId::cntnrSpan')]")  # Ouvrir la liste deroulante
            cliquer("//a[contains(@id, 'AP1:r1:0:SearchResultsTable:r4:0:mapNameId::dropdownPopup::popupsearch')]")  # Rechercher

            mi = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'SearchResultsTable:r4:0:mapNameId::_afrLovInternalQueryId:value00::content')]")))
            mi.send_keys(mappe_importation)  # Inserer creation_article_mobile dans la Mappe d'importation # tnr1_8687
            cliquer("//button[text() = 'Rechercher' and contains(@id, 'mapNameId::_afrLovInternalQueryId::search')]")
            cliquer("//div[contains(@id, 'meId_afrLovInternalTableId::db')]//span[text() = '"+ mappe_importation + "']")  # Rechercher tnr1_8687
            driver.get_screenshot_as_file('C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/taux_dep.png')
            fonctionnel.vers_Word("Renseigner la mappe d'importation",
                                  'C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/taux_dep.png')
            cliquer("//td[contains(@id, 'ResultsTable:r4:0:mapNameId::lovDialogId::_fcc')]//button[text() = 'OK']")  # OK
            # U.Cliquer("//div[contains(@id, 'SearchResultsTable:r4:0:mapNameId::dropdownPopup::dropDownContent::db')]//td/span[text()='creation_article_mobile']")

            Fichier = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'SearchResultsTable:r4:0:if1::content')]")))
            Fichier.send_keys("C:/Users/" + os.getlogin() + "/Desktop/Fichier_importation_tnr/" + csv_nom)  # !!!!!!!!! TEST_544_Crea_masse_ACC_GP.csv
            time.sleep(7)
            # U.Cliquer("//button[contains(@id, 'SearchResultsTable:uploadFileCB')]") CHARGER LE FICHIER sans cliquer sur avoir cliquer sur Verifier la mappe
            cliquer("//button[contains(@id, 'SearchResultsTable:reviewMapCB')]")  # Verifier la mappe
            driver.get_screenshot_as_file('C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/verifier.png')
            time.sleep(8)
            driver.get_screenshot_as_file('C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/verifier1.png')
            cliquer("//div[contains(@id, '0:pt1:AP1:cb2')]")  # Charger le fichier
            time.sleep(3)
            # U.driver.find_element_by_tag_name('body').send_keys(U.Keys.CONTROL + U.Keys.TAB)
            # U.wait.until(U.EC.presence_of_element_located((U.By.XPATH, "//*[text() = 'Mappage de données']")))
            driver.get_screenshot_as_file('C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/encours.png')
            try:
                wait.until(EC.presence_of_element_located((By.XPATH, "//div[text() = 'Confirmation']")))
            finally:
                driver.get_screenshot_as_file('C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/Confirmation.png')
                cliquer("//button[contains(@id, ':0:pt1:AP1:cb11')]")  # OK
            time.sleep(5)

            fonctionnel.vers_Word("Verifier la mappe", 'C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/verifier.png')
            fonctionnel.vers_Word("Description de la verification du mappe", 'C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/verifier1.png')
            fonctionnel.vers_Word("En cours de confirmation", 'C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/encours.png')
            fonctionnel.vers_Word("Confirmation !", 'C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/Confirmation.png')

            cliquer("//img[@alt = 'Actualiser']")
            time.sleep(3)
            driver.get_screenshot_as_file('C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/Resultat.png')
            fonctionnel.vers_Word("Chargement du fichier de mappe d'importation en attente", 'C:/Users/'+os.getlogin()+'/Desktop/TNR-'+str(num_tnr)+'/Resultat.png')

            for i in range(15):
                scroll("//*[text() = 'Colonnes masquées']")
                cliquer("//button[text() = 'Actualiser le total']")
                time.sleep(7)
                try:
                    WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//*[text() = 'Aucun résultat trouvé.']" or "//tr/td[7]/span[@class = 'x2o2']")))
                    while True and i <= 13:
                        print("Aucun resultat trouve. La page va s'actualiser encore une fois")
                        break
                except TimeoutException as e:
                    print("Resultat obtenu")

                    try:
                        WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//img[@title = 'Terminé avec succès']")))
                        while True:
                            print("Succes")
                            driver.get_screenshot_as_file("C:/Users/SXBS5132/Desktop/TNR-" + str(num_tnr) + "/result.png")
                            fonctionnel.vers_Word("Statut de l'import apres " + str(i) + " actualisations du total: Termine avec succes",
                                                  "C:/Users/xxx/Desktop/TNR-" + str(num_tnr) + "/result.png")
                            break
                    except TimeoutException:
                        driver.refresh()
                        time.sleep(5)
                        cliquer("//button[text() = 'Actualiser le total']")
                        time.sleep(5)
                        cliquer("//button[text() = 'Actualiser le total']")
                        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//img[@title = 'Terminé avec succès']")))
                        while True:
                            print("Succes")
                            driver.get_screenshot_as_file("C:/Users/SXBS5132/Desktop/TNR-" + str(num_tnr) + "/result.png")
                            fonctionnel.vers_Word("Statut de l'import apres " + str(i) + " actualisations du total: Termine avec succes",
                                                  "C:/Users/xxx/Desktop/TNR-" + str(num_tnr) + "/result.png")
                            break

                    except TimeoutException as e:
                        print("Echec")
                        fonctionnel.vers_Word("Statut de l'import connu! FAIL",
                                              "C:/Users/xxx/Desktop/TNR-" + str(num_tnr) + "/result.png")
                    break
            #for i in range(15):
            #    cliquer("//img[@alt = 'Actualiser']")
            #    try:
            #        statut_attente = WebDriverWait(driver, 10).until(EC.presence_of_element_located(
            #            (By.XPATH, "//span[contains(@id, ':SearchResultsTable:_ATp:blt:0:soc1::content') and text() = 'En attente']")))
            #        while True:
            #            try:
            #                cliquer("//img[@alt = 'Actualiser']")
            #                time.sleep(5)
            #                print('Statut en attente')
            #                break
            #            except StaleElementReferenceException as e:
            #                print("StaleElementReferenceException")
            #    except TimeoutException as e:
            #         print('Resultat connu')
            #         scroll("//img[@alt = 'Actualiser']")
            #         #table_result = driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:pt1:AP1:r1:0:SearchResultsTable:_ATp:blt::ch::d2::t2')
            #         #for i in range(5):
            #         #   table_result.send_keys(Keys.ARROW_DOWN)
            #         try:
            #            WebDriverWait(driver, 7).until(EC.presence_of_element_located((By.XPATH,
            #                    "//div[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:pt1:AP1:r1:0:SearchResultsTable:_ATp:blt::db']")))
            #            while True:
            #                scroll("//div[@id = '_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:pt1:AP1:r1:0:SearchResultsTable:_ATp:blt::db']")
            #                print("Pas de resultat obtenu !")
            #                break
            #         finally:
            #            driver.get_screenshot_as_file("C:/Users/SXBS5132/Desktop/TNR-" + str(num_tnr) + "/result.png")
            #            fonctionnel.vers_Word("Statut de l'import connu !", "C:/Users/SXBS5132/Desktop/TNR-" + str(num_tnr) + "/result.png")
            #         try:
            #             succes = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//img[@title = 'Terminé avec succès']")))
            #             while True:
            #                 print("Success")
            #                 break
            #             break
            #         except TimeoutException as e:
            #             print("Fail")
            #             break
        except TimeoutException as e:
            print("On ne peut pas ajouter des articles au lot")
            scroll("//td[text() = 'Ajouter des articles au lot']")
            driver.get_screenshot_as_file("C:/Users/xxx/Desktop/TNR-" + str(num_tnr) + "/Ajouter_lot_article_griser.png")
            fonctionnel.vers_Word("On ne peut pas ajouter des articles au lot",
                                  "C:/Users/xxx/Desktop/TNR-" + str(num_tnr) + "/Ajouter_lot_article_griser.png")

    def consulter_les_articles(self, num_article, page_gestion_article_disponible):
        if page_gestion_article_disponible == False:
            cliquer("//a[text() = 'Gestion des informations produit']")
            cliquer("//img[@title = 'Tâches']")
            cliquer("//a[text() = 'Gérer les articles']")
        elif page_gestion_article_disponible == True:
            pass
        Article = wait.until(EC.presence_of_element_located((By.XPATH, "//input[contains(@id, 'region2:0:efqrp:value00::content')]")))
        if Article.get_attribute('value') != '':
            Article.clear()
            Article.send_keys(num_article)
        else:
            Article.send_keys(num_article)
        cliquer("//button[text() = 'Rechercher']")

    def verification_Tx_Depre(self, num_tnr, afrrk, num_article):
        if num_tnr == '86' and afrrk == '2':
            try:
                orga = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH,
                                            "//tr[@_afrrk = '"+str(afrrk)+"']//span[contains(@id, 'effcpOrganizationId5::content')]")))
                print("L'organisation ", orga.text, "est en consideration")
                cliquer("//tr[@_afrrk = '"+str(afrrk)+"']//a[text() = '" + str(num_article) + "']")
            except TimeoutException as e:
                orga = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH,
                                            "//tr[@_afrrk = '0']//span[contains(@id, 'effcpOrganizationId5::content')]")))
                print("L'organisation ", orga.text, "est en consideration")
                cliquer("//tr[@_afrrk = '0']//a[text() = '" + str(num_article) + "']")
        elif num_tnr == '86' and afrrk == '3':
            try:
                orga = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH,
                                            "//tr[@_afrrk = '"+str(afrrk)+"']//span[contains(@id, 'effcpOrganizationId5::content')]")))
                print("L'organisation ", orga.text, "est en consideration")
                cliquer("//tr[@_afrrk = '"+str(afrrk)+"']//a[text() = '" + str(num_article) + "']")
            except TimeoutException as e:
                orga = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH,
                                            "//tr[@_afrrk = '1']//span[contains(@id, 'effcpOrganizationId5::content')]")))
                print("L'organisation ", orga.text, "est en consideration")
                cliquer("//tr[@_afrrk = '1']//a[text() = '" + str(num_article) + "']")
        else:
            orga = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH,
                                        "//tr[@_afrrk = '"+str(afrrk)+"']//span[contains(@id, 'effcpOrganizationId5::content')]")))
            print("L'organisation ", orga.text, "est en consideration")
            cliquer("//tr[@_afrrk = '" + str(afrrk) + "']//a[text() = '" + str(num_article) + "']")


        cliquer("//div[contains(@id, 'pt1:ap1:sspec::ti')]//a[text() = 'Spécifications']")
        cliquer("//a[text() = 'SISTER: Marketing']")
        # scroll("//label[text() = 'Taux de dépréciation du stock']")
        scroll("//div[@title = 'Informations générales']")
        driver.get_screenshot_as_file("C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Tx_Depre.png")
        try:
            tx_dep = driver.find_element_by_xpath("//input[contains(@name, 'GENERAL_INFO_AG:0:XXSCFInvDepreciationRateAT')]")
            # Modifier la valeur selon la derniere mise a jour dans le fichier Excel du tnr 87. Par exemple, remplacer 52 par 58 dans les 4 lignes ci-dessous
            if tx_dep.get_attribute('value') in ['59', '93', '73']:
                print("Le taux de depreciation est reste inchange apres l'import du nouveau fichier de mappe d'importation", tx_dep.get_attribute('value'))
                fonctionnel.vers_Word("Le taux de depreciation est reste inchange apres l'import du nouveau fichier de mappe d'importation FEH_IO_MASTER",
                    "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Tx_Depre.png")
            elif tx_dep.get_attribute('value') not in ['59', '93', '73']:
                print("Le taux de depreciation a ete modifie apres le chargement du mappe d'importation:", tx_dep.get_attribute('value'))
                fonctionnel.vers_Word("Verification de la MAJ du taux de depreciation FEH_IO_MASTER: Modification effectuee",
                    "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Tx_Depre.png")
            scroll("//div[@class = 'x1ij']/a[text() = 'Gérer les articles']")
            # cliquer("//a[2][contains(@id, 'MainAreaTab7::rmAbv') and contains(@title, 'Fermer l')]")
            cliquer("//div[@class = 'x1ij']/a[text() = 'Gérer les articles']")
        except NoSuchElementException as e:
            tx_dep1 = driver.find_element_by_xpath("//span[contains(@id, 'XXSCFInvDepreciationRateAT::content')]")
            if tx_dep1.text in ['59', '93', '73']:
                print("Le taux de depreciation est reste inchange apres l'import du nouveau fichier de mappe d'importation")
                fonctionnel.vers_Word("Le taux de depreciation est reste inchange apres l'import du nouveau fichier de mappe d'importation SCF_OSA_IO_01",
                    "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Tx_Depre.png")
            elif tx_dep1.text not in ['59', '93', '73']:
                print("Le taux de depreciation a ete modifie apres le chargement du mappe d'importation:",
                      tx_dep1.text)
                fonctionnel.vers_Word("Verification de la MAJ du taux de depreciation SCF_OSA_IO_01: Modification effectue",
                    "C:/Users/" + os.getlogin() + "/Desktop/TNR-" + str(num_tnr) + "/Tx_Depre.png")
            scroll("//div[@class = 'x1ij']/a[text() = 'Gérer les articles']")
            # cliquer("//a[2][contains(@id, 'MainAreaTab7::rmAbv') and contains(@title, 'Fermer l')]")
            cliquer("//div[@class = 'x1ij']/a[text() = 'Gérer les articles']")

class fn_tnr90:
    def __init__(self, variable, variable1, r, r1, r2):
        self.variable = variable
        self.variable1 = variable1
        self.r = r
        self.r1 = r1
        self.r2 = r2

    def acceder_aux_dossiers_partages(self):
        Nav = wait.until(EC.element_to_be_clickable((By.XPATH, "//a[@title = 'Navigateur']")))
        try:
            Nav.click()
        except StaleElementReferenceException as e:
            driver.refresh()
            Nav.click()
        except ElementNotInteractableException as e:
            driver.refresh()
            Nav.click()

        cliquer("//div[@title = 'Outils']")
        Etat_analyse = wait.until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'Etats et analyses']")))
        driver.execute_script("arguments[0].click();", Etat_analyse)
        try:
            cliquer("//a[@title = 'Sélecteur de hiérarchie']")
        except TimeoutException as e:
            print("\n NB: IL N'Y A PLUS DE BOUTON QUI S'APPELLE Sélecteur de hiérarchie MAIS Sélecteur hiérarchique \n")
            cliquer("//img[@title = 'Sélecteur hiérarchique']")
        cliquer("//span[text() = 'Dossiers partagés']")
        time.sleep(5)

    def chercher_element_personalise(self):
        table = driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:_FOTsr1:0:pt1:t2::scroller')
        for i in range(50):
            table.send_keys(Keys.ARROW_DOWN)
            try:
                perso = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.XPATH, "//a[@title = 'Personnalisé']")))
                if perso.is_displayed():
                    print("L'élément 'Personnalisé' trouvé !")
                    perso.click()
                    break
            except TimeoutException as e:
                print(i)
                pass

    def Dashboard_Nir(self):
        try:
            time.sleep(5)
            cliquer("//a[contains(@title, 'approvisionnement')]")
        except StaleElementReferenceException as e:
            element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//a[contains(@title, 'approvisionnement')]")))
            driver.execute_script("arguments[0].click();", element)

        time.sleep(5)
        cliquer("//a[contains(text(), 'Product Management')]")
        cliquer("//a[@title = 'SISTER']")
        cliquer("//a[@title = 'Dashboard']")
        cliquer("//span[text() = 'Dashboard NIR']")

    def assurer_conformite_de_la_page(self):
        for i in range(3):
            try:
                Articles = wait.until(EC.presence_of_element_located((By.XPATH, "//a[text() = 'Actualiser']")))
                if Articles.is_displayed():
                    driver.get_screenshot_as_file(Directory + "/Articles.png")
                    fonctionnel.vers_Word("Vérifier que la page s'organise en: Avancement par statut, Retard et Articles", Directory + "/Articles.png")
                    cliquer("//a[text() = 'Retard']")
                    driver.get_screenshot_as_file(Directory + "/retard.png")
                    fonctionnel.vers_Word("", Directory + "/retard.png")
                    break
            except TimeoutException as e:
                print("En attente que la page s'actualise totalement")
                driver.refresh()
        time.sleep(7)

    def direction_bas_articles(self):
        driver.refresh()
        time.sleep(10)
        xpframe = wait.until(EC.frame_to_be_available_and_switch_to_it((By.XPATH, "//iframe[@onload = 'AdfDhtmlInlineFramePeer.__iframeLoadHandler(event)']")))
        time.sleep(7)
        scroll("//*[text() = 'Approuvé' and @pointer-events]")
        time.sleep(10)
        #frameA = driver.find_element_by_xpath("//iframe[@id = 'd:dashboard~p:k6k5brjrd0rr5meb~x:r658ibukq523jhs4']")
        frameA = driver.find_element_by_xpath("//iframe[@id = 'd:dashboard~p:k6k5brjrd0rr5meb~x:h2neggpga11kel98']")
        driver.switch_to.frame(frameA)
        cliquer("//img[@id = 'xdo:viewFormatIcon']")

    def switch_frame(self):
        time.sleep(5)
        try:
            FRAME = driver.find_element_by_xpath("//div[@bindingid='biExecBinding1']/iframe")
            fr = FRAME.get_attribute('name')
            print(fr)
            driver.switch_to.frame(fr)
        except UnexpectedAlertPresentException as e:
            print("UnexpectedAlertPresentException Error")
            pass
            FRAME = driver.find_element_by_xpath("//div[@bindingid='biExecBinding1']/iframe")
            fr = FRAME.get_attribute('name')
            print(fr)
            driver.switch_to.frame(fr)

    def checker_les_3_parties_de_la_page(self):
        global iframe0, iframe1
        seq = driver.find_elements_by_tag_name('iframe')
        print("No of frames present in the web page are: ", len(seq))

        iframe0 = driver.find_elements_by_tag_name('iframe')[0]  # -----------------> iframe0
        print("\nID iframe avec indice 0: ", iframe0.get_attribute('ID'))
        iframe1 = driver.find_elements_by_tag_name('iframe')[1]  # -----------------> iframe1
        print("\nID iframe avec indice 1: ", iframe1.get_attribute('ID'))

        testframe = iframe1.get_attribute('ID')
        print("//iframe[@name = '" + str(testframe) + "']")

        frames = driver.find_elements_by_tag_name("iframe")
        print("\nLe nombre d'iframes present actuellement est: ", len(frames), "avec: ")
        for f in frames:
            print("\nID: ", f.get_attribute('id'), "et NAME:", f.get_attribute('name'))
        time.sleep(8)
        driver.switch_to.frame(iframe1)

        iframexp1 = driver.find_element_by_xpath("//iframe[@name = 'd:dashboard~p:k6k5brjrd0rr5meb~x:40hmetbh9d802nhb']")
        driver.switch_to.frame(iframexp1)
        ### try:
        ###     iframexp1 = driver.find_element_by_xpath("//iframe[@name = 'd:dashboard~p:k6k5brjrd0rr5meb~x:182g4b2f21km98oa']")
        ###     driver.switch_to.frame(iframexp1)
        ### except NoSuchElementException as e:
        ###     time.sleep(2)
        ###     #print("//iframe[@name = '"+str(iframe1.get_attribute('ID'))+"']")
        ###     #iframe1 = driver.find_elements_by_tag_name('iframe')[1]
        ###     testframe = iframe1.get_attribute('ID')
        ###     print("\n NB LE IFRAME A CHANGE et est: !", testframe, "\n")
        ###     #iframexp1 = driver.find_element_by_xpath("//iframe[@name = '"+str(testframe)+"']")
        ###     driver.switch_to.frame(testframe)

        time.sleep(3)
        cliquer("//a[text() = 'Retard']")
        time.sleep(15)
        driver.save_screenshot(Directory + "/statut_retard.png")
        fonctionnel.vers_Word("Vérifier que la page s'organise en: (1)Avancement par statut (2)Retard", Directory + "/statut_retard.png")

        fn_tnr90.direction_bas_articles(self)
        time.sleep(7)
        driver.save_screenshot(Directory + "/statut_Article.png")
        fonctionnel.vers_Word("Vérifier que la page s'organise en: (3) Bas : Articles", Directory + "/statut_Article.png")
        time.sleep(10)

    def conversion_et_transfert_du_fichier_xlsx(self):
        # Conversion du fichier xlsx en csv et transfert du fichier dans le dossier TNR-90
        cliquer1("//li[@fmid = '3']/div/a[@class = 'masterMenuItem item']//following::div[text() = 'Excel (*.xlsx)'][2]")
        Direc = "C:/Users/" + os.getlogin()
        Direct = Direc + "/Downloads/NIR Count12 Report_Détail de l'encours.xlsx"
        if os.path.exists(Direc):
            print('File exist')
            list_of_files = glob.glob("C:/Users/" + os.getlogin() + "/Downloads/*.xlsx")
            latest_file = max(list_of_files, key=os.path.getmtime)
            print(latest_file)
        else:
            print('File not found')

        separator = os.path.normpath("/")
        if separator != "/":
            path = re.sub(re.escape(separator), "/", latest_file)
        print("Nouveau chemin xlsx\n", path)
        read_file = pd.read_excel(path)
        print("En format xlsx:\n", read_file)
        read_file.to_csv(Directory + "/NIR Count12 ReportOriginal.csv", index=None, header=False)

    def choisir_mois_concerne(self):
        driver.refresh()
        fn_tnr90.switch_frame(self)
        try:
            cliquer("//select[contains(@name, 'd_0')]")
        except TimeoutException as e:
            print("TimeoutException")
            cliquer("//select[contains(@name, '6_0')]")

        fn_tnr90.choisir_mois_recent('d_0', '6_0')
        time.sleep(5)
        driver.switch_to.default_content()

    def choisir_mois_recent(variable = '', variable1 = ''):
        def section2_mois(variable = '', variable1 = ''):
            global Mois_option
            Mois_option = WebDriverWait(driver, 15).until(EC.presence_of_all_elements_located((By.XPATH, "//td/select[contains(@name, '" + str(variable) + "')]/option")))
            print("\nLe nombre de mois disponible: ", len(Mois_option))
            Mois_recent = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, "//td/select[contains(@name, '" + str(variable) + "')]/option[" + str(len(Mois_option)) + "]")))
            print("\nLe dernier mois est: ", Mois_recent.text)
        try:
            section2_mois(variable)
        except TimeoutException as e:
            section2_mois(variable1)

        def mois_recent(variable = '', variable1 = ''):
            cliquer("//select[contains(@name, '" + str(variable) + "')]")
            time.sleep(2)
            scroll("//td/select[contains(@name, '" + str(variable) + "')]/option[" + str(z) + "]")
            cliquer("//td/select[contains(@name, '" + str(variable) + "')]/option[" + str(z) + "]")
            time.sleep(3)
        for z in range(len(Mois_option) - 1, len(Mois_option) + 1):
            print('Indice du mois:', z)
            # util.scroll("//td/select[contains(@name, '"+str(variable)+"')]/option[" + str(len(Mois_option)) + "]")
            try:
                mois_recent(variable)
            except TimeoutException as e:
                print("\nTimeoutException\n")
                mois_recent(variable1)

    def semaine_details(r):
        Semaine_option = WebDriverWait(driver, 20).until(EC.presence_of_all_elements_located((By.XPATH, "//td/select[contains(@name, '" + str(r) + "')]/option")))
        print("Le nombre de date disponible: ", len(Semaine_option))
        Semaine_recent = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, "//td/select[contains(@name, '" + str(r) + "')]/option[" + str(len(Semaine_option)) + "]")))
        print("La derniere semaine est: ", Semaine_recent.text)
        cliquer("//td/select[contains(@name, '" + str(r) + "')]/option[" + str(len(Semaine_option)) + "]")

    def semaine_du_calendrier(r1, r2):
        time.sleep(5)
        try:
            D_S = driver.find_element_by_xpath("//td/select[contains(@name, 'd_1')]")
            D_S.click()
            fn_tnr90.semaine_details(r1)
        except NoSuchElementException as e:
            cliquer("//td/select[contains(@name, '_6_1')]")
            time.sleep(3)
            fn_tnr90.semaine_details(r2)
        time.sleep(5)

    def checker_mois_et_fin_de_semaine_actuelle(self):
        driver.switch_to.default_content()
        fn_tnr90.switch_frame(self)
        try:
            cliquer("//select[contains(@name, 'c_0')]")
        except TimeoutException as e:
            print("\nTimeoutException\n")
            cliquer("//select[contains(@name, '_11_0')]")

        fn_tnr90.choisir_mois_recent('c_0', '_11_0')
        driver.save_screenshot(Directory + "/recent_date.png")
        fonctionnel.vers_Word("Vérifier que les données sont à jour pour le mois actuel et la date de fin de semaine", Directory + "/recent_date.png")
        # scroll("//*[text() = 'Approuvé' and @transform = 'matrix(1,0,0,1,0,0)']")

    def date_retard_recent(self):
        #iframexp1 = driver.find_element_by_xpath("//iframe[@name = 'd:dashboard~p:k6k5brjrd0rr5meb~x:182g4b2f21km98oa']")
        iframexp1 = driver.find_element_by_xpath("//iframe[@name = 'd:dashboard~p:k6k5brjrd0rr5meb~x:40hmetbh9d802nhb']")
        driver.switch_to.frame(iframexp1)
        time.sleep(3)
        cliquer("//a[text() = 'Retard']")
        frameint = driver.find_element_by_xpath("//iframe[@name = 'xdo:docframe1']")

        driver.switch_to.frame(frameint)
        cliquer("//td[@class = 'tableHeaderCell tableLov' and @id = '35']")  # filtre
        cliquer("//a[@class = 'choice']/input[@name = 'filterall']")  # unselect all

        dates = WebDriverWait(driver, 15).until(EC.presence_of_all_elements_located((By.XPATH, "//li/a[@class = 'choice' and contains(@id, 'ev')]")))
        print("Le nombre de dates disponible est ", len(dates))
        cliquer("//li[" + str(len(dates)) + "]/a[@class = 'choice' and contains(@id, 'ev')]")
        cliquer("//input[@value = 'OK']")

    def checker_semaine_attendue_actuelle(self):
        time.sleep(10)
        # util.driver.switch_to.frame(iframe1)
        fn_tnr90.date_retard_recent(self)
        time.sleep(15)
        driver.save_screenshot(Directory + "/retard_recent.png")
        fonctionnel.vers_Word("Vérifier que les données 'Retard' sont mises à jour pour la semaine attendue la plus récente", Directory + "/retard_recent.png")

class fn_tnr94:
    def __init__(self, classe):
        self.classe = classe

    def test(self, classe):
        #try:
        #    cliquer("//*[@aria-label = 'Navigateur']")
        #    scroll("//span[text() = 'Mon entreprise']")
        #    cliquer1("//span[text() = 'Mon entreprise']")
        #    cliquer("//span[text() = 'Configuration et maintenance']")
        #    cliquer("//a[@id = 'pt1:r1:0:r0:0:r1:0:AP1:soc2::drop']")
        #except TimeoutException as e:
        #    cliquer("//div[@id = 'itemNode_MyEnterprise_setup_and_maintenance']")
        #    cliquer("//a[@id = 'pt1:r1:0:r0:0:r1:0:AP1:soc2::drop']")
        cliquer("//*[@aria-label = 'Navigateur']")
        print("yes")
        cliquer("//span[text() = 'Accueil']")
        cliquer("//div[@id = 'itemNode_MyEnterprise_setup_and_maintenance']")
        cliquer("//a[@id = 'pt1:r1:0:r0:0:r1:0:AP1:soc2::drop']")
        cliquer("//li[text() = 'Product Management']")
        cliquer("//td[text() = 'Articles']")
        cliquer("//a[contains(text(), 'Gérer les classes d')]")
        for i in range(3):
            cliquer("//a[@title = 'Développer']")
            time.sleep(3)
        cliquer("//tr[@_afrrk = '13']//a[@title = 'Développer']")
        cliquer("//a[text() = '"+str(classe)+"']")
        cliquer("//a[@id = 'pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:sdi7::disAcr' and text() = 'Modèles et formats']")
        # COMMERCIAL, FTT MOBILES 6 SEMAINES, Fini, RECONDITIONNE GP,
        time.sleep(7)

        def controle_des_matrices(Nom):
            time.sleep(5)
            scroll("//a[text() = 'SISTER: Direction commerciale']")
            time.sleep(3)

            # cliquer("//a[text() = 'SISTER: Marketing']")
            cliquer("//a[text() = 'SISTER: Marketing']")
            time.sleep(7)
            scroll("//h3[text() = 'Informations générales']")
            driver.save_screenshot(Directory + "/controle_marketing1.png")
            fonctionnel.vers_Word(str(Nom) + " - MARKETING : Informations générales et Génération EAN",
                                  Directory + "/controle_marketing1.png")

            scroll("//h3[text() = 'Génération EAN']")
            time.sleep(1)
            driver.save_screenshot(Directory + "/controle_marketing2.png")
            fonctionnel.vers_Word("Controle: Référence fournisseur",
                                  Directory + "/controle_marketing2.png")

            try:
                scroll("//h3[text() = 'Référence fournisseur']")
                driver.save_screenshot(Directory + "/controle_marketing3.png")
                fonctionnel.vers_Word("Controle: Traçabilité",
                                      Directory + "/controle_marketing3.png")
            except TimeoutException as e:
                pass

            scroll("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
            cliquer("//a[text() = 'SISTER: Direction commerciale']")
            time.sleep(5)
            driver.save_screenshot(Directory + "/Dirco.png")
            fonctionnel.vers_Word(str(Nom) + " - Direction commerciale : Cluster", Directory + "/Dirco.png")

            cliquer("//a[text() = 'SISTER: Achat']")
            time.sleep(5)
            driver.save_screenshot(Directory + "/Achat.png")
            fonctionnel.vers_Word(str(Nom) + " - ACHAT : Informations achat et Référence fournisseur",
                                  Directory + "/Achat.png")

            cliquer("//a[text() = 'SISTER: Pilote']")
            time.sleep(5)
            scroll("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
            driver.save_screenshot(Directory + "/Pilote.png")
            fonctionnel.vers_Word(str(Nom) + " - Pilote : Gestion Produit", Directory + "/Pilote.png")

            scroll("//h3[text() = 'EAN']")
            driver.save_screenshot(Directory + "/Pilote1.png")
            fonctionnel.vers_Word("Controle: Référence fournisseur", Directory + "/Pilote1.png")

            scroll("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
            cliquer("//a[text() = 'SISTER: Référentiel']")
            scroll("//a[text() = 'SISTER: Référentiel']")
            time.sleep(5)

            driver.save_screenshot(Directory + "/Referentiel.png")
            fonctionnel.vers_Word(str(Nom) + " - Référentiel : Génération EAN", Directory + "/Referentiel.png")

            try:
                scroll("//h3[text() = 'Familles Supply Chain']")
                driver.save_screenshot(Directory + "/supplychain.png")
                fonctionnel.vers_Word("Controle: Familles Supply Chain", Directory + "/supplychain.png")
            except TimeoutException as e:
                print("Pas de Familles Supply Chain")

        if classe == 'MOBILE GP':
            #def controle():
            #    ref = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, "//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '"+str(i)+"']//td[1]//span/span/span")))
            #    if ref.text == 'COMMERCIAL':
            #        print(ref.text)
            #        cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '"+str(i)+"']//td[1]//span/span/span")
            #        cliquer("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
#
            #        controle_des_matrices()
#
            #    elif ref.text == 'Fini':
            #        print(ref.text)
            #        cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '"+str(i)+"']//td[1]//span/span/span")
            #        cliquer("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
            #        controle_des_matrices()
#
            #    elif ref.text == 'RECONDITIONNE GP':
            #        print(ref.text)
            #        cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '"+str(i)+"']//td[1]//span/span/span")
            #        cliquer("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
            #        controle_des_matrices()
#
            #    elif ref.text == 'SAV MOBILE':
            #        print(ref.text)
            #        cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(i) + "']//td[1]//span/span/span")
            #        cliquer("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
            #        controle_des_matrices()
            #    else:
            #        print("Reference pas trouve")

            ### for i in range(10):
            ###     try:
            ###         # doit modifier le ref et inclure startrow dedans !!!!!!!!!
###
            ###         ref = WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.XPATH, "//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '"+str(i)+"']//td[1]//span/span/span")))
###
###
            ###         #table_matrice.send_keys(Keys.ARROW_DOWN)
            ###         if ref.text != '':
            ###             print(ref.text, i)
            ###             table_matrice.send_keys(Keys.ARROW_DOWN)
            ###             time.sleep(1)
            ###         elif ref.text == '':
            ###             for k in range(4):
            ###                 table_matrice.send_keys(Keys.ARROW_DOWN)
            ###             print("Derniere valeur par lot de 6 avec startrow: ", i - 5, ref.text, i, "- Fin du lot")
            ###             time.sleep(5)
###
            ###     except TimeoutException as e:
            ###         print("Fin premiere partie")
            ###         break

                #try:
                #    #scroll("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '"+str(i+1)+"']//td[1]//span/span/span")

                #if i % 6 == 0 and i > 0:
                    #print(i+1)
                    #table_matrice.send_keys(Keys.PAGE_DOWN)
                    #scroll("//table[contains(@summary, 'Modèles d') and @_startrow = '"+str(i+1)+"'][1]//tr[@_afrrk = '"+str(i+1)+"']//td[1]//span/span/span")

            #for m in range(11, 100, 6):
            table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')

            m = 5
            while True:
                print(m)
                for p in range(m-5, m+2):
                    try:
                        ref = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.XPATH,
                        "//table[contains(@summary, 'Modèles d') and @_startrow = '" + str(m-5) + "'][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")))
                        if ref.text != '':
                            print(ref.text, p, "----------> *")
                            if ref.text == 'COMMERCIAL':
                                print(ref.text)
                                cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")
                                driver.save_screenshot(Directory + "/commercial.png")
                                fonctionnel.vers_Word("Cliquer sur Commercial", Directory + "/commercial.png")
                                cliquer1("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")

                                controle_des_matrices('COMMERCIAL')
                                time.sleep(2)
                                htmlelement = driver.find_element_by_tag_name('html')
                                htmlelement.send_keys(Keys.HOME)

                            elif ref.text == 'Fini':
                                print(ref.text, p, "----------> *")
                                #cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")
                                cliquer("//span[text() = 'Fini']")
                                #cliquer1("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
                                driver.save_screenshot(Directory + "/fini.png")
                                fonctionnel.vers_Word("Cliquer sur Fini", Directory + "/fini.png")
                                controle_des_matrices('Fini')
                                time.sleep(2)
                                htmlelement = driver.find_element_by_tag_name('html')
                                htmlelement.send_keys(Keys.HOME)
                                table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                                for z in range(5):
                                    table_mat.send_keys(Keys.ARROW_DOWN)

                            elif ref.text == 'RECONDITIONNE GP':
                                print(ref.text, p, "----------> *")
                                #cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")
                                cliquer("//span[contains(@id, 'inputText5') and text() = 'RECONDITIONNE GP']")
                                #cliquer1("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
                                driver.save_screenshot(Directory + "/reconditionne_gp.png")
                                fonctionnel.vers_Word("Cliquer sur Reconditionne GP", Directory + "/reconditionne_gp.png")
                                controle_des_matrices('RECONDITIONNE GP')
                                time.sleep(2)
                                htmlelement = driver.find_element_by_tag_name('html')
                                htmlelement.send_keys(Keys.HOME)
                                time.sleep(2)
                                table_mat2 = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                                for z in range(5):
                                    table_mat2.send_keys(Keys.ARROW_DOWN)

                            elif ref.text == 'SAV MOBILE':
                                print(ref.text, p, "----------> *")

                                cliquer("//span[contains(@id, 'inputText5') and text() = 'RECONDITIONNE GP']")
                                #cliquer1("//span[contains(@id, 'inputText5') and text() = 'RECONDITIONNE GP']")
                                for x in range(5):
                                    time.sleep(3)
                                    table_mat.send_keys(Keys.ARROW_DOWN)
                                cliquer("//span[text() = 'SAV MOBILE']")
                                #cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")
                                #table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                                #for z in range(5):
                                #    table_mat.send_keys(Keys.ARROW_DOWN)
                                #    sav = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.XPATH, "//span[text() = 'SAV MOBILE']")))
                                #    if sav.is_displayed():
                                #        cliquer1("//span[text() = 'SAV MOBILE']")
                                #        break
                                #    else:
                                #        pass

                                #cliquer("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
                                driver.save_screenshot(Directory + "/sav_mobile.png")
                                fonctionnel.vers_Word("Cliquer sur SAV mobile", Directory + "/sav_mobile.png")
                                controle_des_matrices('SAV MOBILE')

                            else:
                                print("Reference pas trouve")

                            time.sleep(2)
                            # ! table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                            time.sleep(2)
                            try:
                                table_mat.send_keys(Keys.ARROW_DOWN)
                            except StaleElementReferenceException as e:
                                time.sleep(2)
                                htmlelement = driver.find_element_by_tag_name('html')
                                # htmlelement.send_keys(Keys.END)
                                htmlelement.send_keys(Keys.HOME)
                                print("StaleElementReferenceException")
                                table_mat1 = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                                table_mat1.send_keys(Keys.ARROW_DOWN)

                        elif ref.text == '':
                            table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                            for k in range(3):
                                table_mat.send_keys(Keys.ARROW_DOWN)
                                time.sleep(5)
                            print(ref.text, p)
                    except TimeoutException as e:
                        table_mat1 = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                        table_mat1.send_keys(Keys.ARROW_DOWN)
                m += 6
                print("new", m)
                try:
                    refcount = len(WebDriverWait(driver, 2).until(EC.presence_of_all_elements_located((By.XPATH,
                    "//table[contains(@summary, 'Modèles d') and @_startrow = '" + str(m - 5) + "'][1]//td[1]//span/span/span"))))
                    if refcount == 6:
                        print("Refcount", refcount)
                        continue
                    elif refcount < 6:
                        continue
                    else:
                        break
                    #elif refcount < 6:
                    #    print(refcount)
                    #    for i in range(refcount + 1):
                    #        table_matrice.send_keys(Keys.ARROW_DOWN)
                    #        time.sleep(1)
                    #        ref = WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.XPATH,
                    #            "//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(i) + "']//td[1]//span/span/span")))
                    #        # table_matrice.send_keys(Keys.ARROW_DOWN)
                    #        print(ref.text, i)
                    #    print("Derniere valeur par lot de 6 avec startrow: ", m - 5, ref.text, i, "- Fin du lot")
                    #    break
                except TimeoutException as e:
                    print(m-5)
                    if refcount < 6:
                        cliquer("//span[text() = 'SAV MOBILE']")
                        driver.save_screenshot(Directory + "/sav_mobile.png")
                        fonctionnel.vers_Word("Cliquer sur SAV mobile", Directory + "/sav_mobile.png")
                        controle_des_matrices('SAV MOBILE')
                        break
                    else:
                        break

                #if ref.text == 'COMMERCIAL':
                #    print(ref.text)
                #    i += 1
                #elif ref.text == 'Fini':
                #    print(ref.text)
                #    i += 1
                #else:
                #    pass

                #ref = driver.find_element_by_xpath("//table[contains(@summary, 'Modèles d')][1]//td[1]/span/span/span")
#
                #if (ref.text in ['COMMERCIAL', 'Fini', 'RECONDITIONE GP', 'SAV MOBILE']) and i < 100:
                #    print(ref.text)
                #else:
                #    print(ref.text)
                #    table_matrice.send_keys(Keys.DOWN)

        elif classe == 'MULTIMEDIA': # startrow = 0, 3, 9,

            table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')

            m = 5
            while True:
                print(m)
                for p in range(m - 5, m + 1):
                    try:
                        ref = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.XPATH,
                        "//table[contains(@summary, 'Modèles d') and @_startrow = '" + str(m - 5) + "'][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")))
                        if ref.text != '':
                            if ref.text == 'Fini':
                                print(ref.text, p, "----------> *")
                                # cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")
                                cliquer1("//span[text() = 'Fini']")
                                # cliquer1("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
                                driver.save_screenshot(Directory + "/fini_multi.png")
                                fonctionnel.vers_Word("Cliquer sur Fini", Directory + "/fini_multi.png")
                                controle_des_matrices('Fini')
                                time.sleep(2)
                                htmlelement = driver.find_element_by_tag_name('html')
                                htmlelement.send_keys(Keys.HOME)
                                table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                                for z in range(2):
                                    table_mat.send_keys(Keys.ARROW_DOWN)

                            elif ref.text == 'RECYCLE':
                                print(ref.text, p, "----------> *")
                                # cliquer("//table[contains(@summary, 'Modèles d')][1]//tr[@_afrrk = '" + str(p) + "']//td[1]//span/span/span")
                                cliquer("//span[text() = 'RECYCLE']")
                                # cliquer1("//div[contains(@_afrptkey, 'showDetailItem3')]//a[text() = 'Spécifications']")
                                driver.save_screenshot(Directory + "/recycle.png")
                                fonctionnel.vers_Word("Cliquer sur RECYCLE", Directory + "/recycle.png")
                                controle_des_matrices('RECYCLE')
                                time.sleep(2)

                            else:
                                print("Reference pas trouve")

                            time.sleep(2)
                            # ! table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                            try:
                                table_mat.send_keys(Keys.ARROW_DOWN)
                            except StaleElementReferenceException as e:
                                time.sleep(2)
                                htmlelement = driver.find_element_by_tag_name('html')
                                # htmlelement.send_keys(Keys.END)
                                htmlelement.send_keys(Keys.HOME)
                                print("StaleElementReferenceException")
                                table_mat1 = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                                table_mat1.send_keys(Keys.ARROW_DOWN)

                        elif ref.text == '':
                            table_mat = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                            for k in range(3):
                                table_mat.send_keys(Keys.ARROW_DOWN)
                                time.sleep(5)
                            print(ref.text, p)
                    except TimeoutException as e:
                        table_mat1 = driver.find_element_by_id('pt1:r1:0:rt:1:r2:0:dynamicRegion1:1:pgt1:ap1:ifr421:1:AT1:_ATp:table1::scroller')
                        table_mat1.send_keys(Keys.ARROW_DOWN)
                m += 6
                print("new", m)
                try:
                    refcount = len(WebDriverWait(driver, 2).until(EC.presence_of_all_elements_located((By.XPATH,
                    "//table[contains(@summary, 'Modèles d') and @_startrow = '" + str(m - 5) + "'][1]//td[1]//span/span/span"))))
                    if refcount == 6:
                        print("Refcount", refcount)
                        continue
                    elif refcount < 6:
                        continue
                    else:
                        break
                except TimeoutException as e:
                    print("Fin.")

    def ajouter_date_de_creation(self):
        cliquer("//button[text() = 'Ajouter des champs']")
        cliquer("//span[text() = 'Date de création']")
        cliquer("//button[text() = 'Ajouter']")
        cliquer("//button[contains(@id, 'addFieldsDialog::ok')]")

    def alimenter_les_champs(self):
        remplir_champ("//input[contains(@id, 'region2:0:efqrp:value00::content')]", '2')
        time.sleep(5)
        print("ici")
        try:
            cliquer("//select[contains(@name, 'Panel1:region2:0:efqrp:operator3')]")
        except Exception as e:
            cliquer1("//select[contains(@name, 'Panel1:region2:0:efqrp:operator3')]")

        #date = driver.find_element_by_xpath("//option[@value = 'Egal ou postérieur à']")

        from selenium.webdriver.support.ui import Select

        select = Select(driver.find_element_by_id('_FOpt1:_FOr1:0:_FONSr2:0:MAt2:0:pt1:ItemC1:0:simplePanel1:region2:0:efqrp:operator3::content'))
        for o in select.options:
            print(o.text)
            if o.text == 'Egal ou postérieur à':
                print(o)
                o.click()
        #actions = ActionChains(driver)
        #actions.move_to_element(date)
        #actions.click(date).perform()
